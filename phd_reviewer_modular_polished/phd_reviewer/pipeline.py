#!/usr/bin/env python3
r"""
Production-oriented multi-agent DOCX reviewer for PhD papers.
v2
Modular edition: agent logic is loaded from phd_reviewer/agents/*.py so each agent can be edited independently.

What this version adds
----------------------
1. Keeps the original .docx as the source of truth.
2. Extracts paragraphs + embedded figures into a structured JSON payload.
3. Runs multiple specialized agents under a chief editor:
   - GrammarAgent
   - ScienceAgent
   - MethodsAgent
   - ResultsAgent
   - CitationStyleAgent
   - Deterministic CrossrefReferenceAgent (DOI / Crossref verification)
4. Supports multi-round review.
5. Applies tracked changes to the original DOCX when possible.
6. Adds real Word comments (bubbles) into the DOCX.
7. Supports comment-only mode for safer scientific review.
8. Adds logging, retries, configurable chunking, optional parallel chunk processing,
   and JSON/YAML config loading.

Recommended usage with xAI / Grok
---------------------------------
.env example:
    XAI_API_KEY=...
    XAI_BASE_URL=https://api.x.ai/v1
    REVIEW_MODEL=grok-4-1-fast-reasoning

PowerShell:
    python -u .\multi_agent_paper_reviewer_v3.py `
      --input .\draft_paper.docx `
      --output .\reviewed_paper.docx `
      --payload-json .\review_payload.json `
      --decisions-json .\review_decisions.json `
      --apply-report-json .\review_apply_report.json `
      --crossref-report-json .\review_crossref_report.json `
      --log-file .\review.log `
      --model grok-4-1-fast-reasoning `
      --rounds 2 `
      --mode edit_and_comment
"""

from __future__ import annotations

import argparse
import base64
import concurrent.futures
import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import textwrap
import time
import unicodedata
import zipfile
import xml.etree.ElementTree as ET
from html import unescape
from io import BytesIO
from collections import defaultdict
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence, Tuple

import requests
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx_revisions import RevisionDocument, RevisionParagraph

try:
    import yaml  # type: ignore
except Exception:  # pragma: no cover - optional import, handled later
    yaml = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover - optional import, handled later
    PdfReader = None


# -----------------------------------------------------------------------------
# Configuration defaults
# -----------------------------------------------------------------------------

load_dotenv()

DEFAULT_MODEL = os.getenv("REVIEW_MODEL", "grok-4-1-fast-reasoning")
DEFAULT_AUTHOR = os.getenv("REVIEW_AUTHOR", "AI Reviewer")
DEFAULT_BASE_URL = os.getenv("OPENAI_BASE_URL") or os.getenv("XAI_BASE_URL") or None
DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("XAI_API_KEY")
DEFAULT_LOG_LEVEL = os.getenv("REVIEW_LOG_LEVEL", "INFO")


def default_enabled_agents() -> List[str]:
    return [
        "overall_review_agent",
        "grammar_agent",
        "science_agent",
        "methods_agent",
        "results_agent",
        "citation_style_agent",
        "literature_review_agent",
        "fact_check_agent",
        "reference_verification",
    ]



# -----------------------------------------------------------------------------
# Data models
# -----------------------------------------------------------------------------

@dataclass
class ParagraphRecord:
    paragraph_id: str
    order: int
    section: str
    section_bucket: str
    style: str
    text: str
    text_hash: str
    in_table: bool
    contains_citation: bool = False


@dataclass
class FigureRecord:
    figure_id: str
    filename: str
    extracted_path: str
    mime_type: str
    caption_guess: str
    caption_paragraph_id: str = ""
    vision_summary: str = ""


@dataclass
class EditSuggestion:
    paragraph_id: str
    find: str = ""
    replace: str = ""
    context_before: str = ""
    context_after: str = ""
    comment: str = ""
    kind: str = "grammar"
    confidence: str = "high"
    source_agent: str = ""
    round_index: int = 1
    replace_entire_paragraph: bool = False


@dataclass
class CommentSuggestion:
    paragraph_id: str = ""
    anchor_text: str = ""
    comment: str = ""
    kind: str = "science"
    source_agent: str = ""
    round_index: int = 1
    figure_id: str = ""


@dataclass
class ReferenceCheckFinding:
    paragraph_id: str
    reference_text: str
    found_doi: str = ""
    crossref_doi: str = ""
    crossref_title: str = ""
    crossref_year: str = ""
    confidence: str = "low"
    issues: List[str] = field(default_factory=list)
    matched: bool = False
    lookup_method: str = ""


@dataclass
class ReviewConfig:
    input: str
    output: str
    payload_json: str = "review_payload.json"
    decisions_json: str = "review_decisions.json"
    apply_report_json: str = "review_apply_report.json"
    crossref_report_json: str = "review_crossref_report.json"
    model: str = DEFAULT_MODEL
    author: str = DEFAULT_AUTHOR
    base_url: Optional[str] = DEFAULT_BASE_URL
    rounds: int = 2
    work_dir: str = "review_workdir"
    vision: bool = False
    mode: str = "edit_and_comment"  # edit_and_comment | comments_only
    chunk_chars: int = 10000
    max_workers: int = 2
    request_timeout: int = 180
    max_retries: int = 3
    crossref_enabled: bool = True
    crossref_mailto: str = ""
    max_references_to_verify: int = 400
    max_fact_check_paragraphs: int = 60
    editing_aggressiveness: str = "substantive"  # conservative | balanced | substantive
    structure_review_enabled: bool = True
    fact_check_enabled: bool = True
    reference_workers: int = 6
    unpaywall_email: str = os.getenv("UNPAYWALL_EMAIL", "")
    open_access_fulltext_enabled: bool = True
    max_oa_lookups: int = 120
    max_fulltext_chars: int = 20000
    max_evidence_snippets: int = 5
    max_comments_per_paragraph: int = 3
    enabled_agents: List[str] = field(default_factory=default_enabled_agents)
    log_file: str = "review.log"
    log_level: str = DEFAULT_LOG_LEVEL


# -----------------------------------------------------------------------------
# Logging
# -----------------------------------------------------------------------------

logger = logging.getLogger("multi_agent_reviewer")


def setup_logging(log_file: str, level: str = "INFO") -> None:
    logger.handlers.clear()
    logger.setLevel(getattr(logging, level.upper(), logging.INFO))
    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")

    console = logging.StreamHandler()
    console.setFormatter(formatter)
    logger.addHandler(console)

    if log_file:
        fh = logging.FileHandler(log_file, encoding="utf-8")
        fh.setFormatter(formatter)
        logger.addHandler(fh)


# -----------------------------------------------------------------------------
# Generic helpers
# -----------------------------------------------------------------------------

CITATION_PATTERNS = [
    re.compile(r"\(([A-Z][A-Za-z'`\-]+(?:\s+et\s+al\.)?(?:,?\s*&\s*[A-Z][A-Za-z'`\-]+)?),\s*(19|20)\d{2}[a-z]?[^)]*\)"),
    re.compile(r"([A-Z][A-Za-z'`\-]+(?:\s+et\s+al\.)?)\s*\((19|20)\d{2}[a-z]?\)"),
]
DOI_RE = re.compile(r"\b(10\.\d{4,9}/[-._;()/:A-Z0-9]+)\b", re.I)
YEAR_RE = re.compile(r"\b(19|20)\d{2}[a-z]?\b")


def text_hash(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:12]


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def normalize_text_loose(text: str) -> str:
    txt = unicodedata.normalize("NFKD", text or "")
    txt = "".join(ch for ch in txt if not unicodedata.combining(ch))
    txt = txt.replace("＆", "&").replace("–", "-").replace("—", "-").replace("−", "-")
    txt = txt.replace("’", "'").replace("`", "'")
    txt = re.sub(r"\bet\s+al\b", "et al", txt, flags=re.I)
    return normalize_ws(txt)


def normalize_author_token(text: str) -> str:
    txt = normalize_text_loose(text).lower()
    txt = re.sub(r"[^a-z0-9'\-]+", " ", txt)
    return normalize_ws(txt)


def contains_in_text_citation(text: str) -> bool:
    txt = text or ""
    return any(p.search(txt) for p in CITATION_PATTERNS)


def first_nonempty(items: Sequence[Optional[str]]) -> str:
    for item in items:
        if isinstance(item, str) and normalize_ws(item):
            return item
    return ""


def split_reference_author_block(author_block: str) -> List[str]:
    txt = normalize_text_loose(author_block)
    txt = re.sub(r"\bet al\.?", "", txt, flags=re.I)
    parts = re.split(r"\s*(?:,\s*|\s*&\s*|\sand\s*)", txt)
    out: List[str] = []
    for part in parts:
        part = normalize_ws(part)
        if not part:
            continue
        surname = re.split(r"\s+", part)[0]
        surname = normalize_author_token(surname)
        if surname:
            out.append(surname)
    return out


def extract_comment_citation_signatures(text: str) -> List[Tuple[str, str]]:
    signatures: List[Tuple[str, str]] = []
    txt = normalize_text_loose(text)
    patterns = [
        re.compile(r"([A-Z][A-Za-z'\-]+(?:\s+et\s+al\.?)?(?:\s*&\s*[A-Z][A-Za-z'\-]+)?)\s*\((19|20)\d{2}[a-z]?\)"),
        re.compile(r"([A-Z][A-Za-z'\-]+(?:\s+et\s+al\.?)?(?:\s*&\s*[A-Z][A-Za-z'\-]+)?),\s*((?:19|20)\d{2}[a-z]?)"),
    ]
    for pattern in patterns:
        for m in pattern.finditer(txt):
            author_part = m.group(1)
            year_text = m.group(2)
            surnames = split_reference_author_block(author_part)
            year_match = re.search(r"(19|20)\d{2}", year_text)
            if surnames and year_match:
                signatures.append((surnames[0], year_match.group(0)))
    return signatures


def safe_initials(author: str) -> str:
    cleaned = re.sub(r"[^A-Za-z]", "", author or "AI")
    return (cleaned[:2] or "AI").upper()


def try_read_json_or_yaml(path: Optional[str]) -> Dict[str, Any]:
    if not path:
        return {}
    config_path = Path(path)
    if not config_path.exists():
        raise FileNotFoundError(f"Config file not found: {config_path}")

    raw = config_path.read_text(encoding="utf-8")
    suffix = config_path.suffix.lower()
    if suffix in {".yaml", ".yml"}:
        if yaml is None:
            raise RuntimeError("PyYAML is not installed. Install pyyaml or use JSON config.")
        data = yaml.safe_load(raw) or {}
    else:
        data = json.loads(raw or "{}")

    if not isinstance(data, dict):
        raise ValueError("Config file must contain a top-level object / mapping.")
    return data


def read_json_cache(path: Path) -> Dict[str, Any]:
    try:
        if path.exists():
            raw = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(raw, dict):
                return raw
    except Exception:
        pass
    return {}


def write_json_cache(path: Path, data: Dict[str, Any]) -> None:
    try:
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def strip_xml_or_html(text: str) -> str:
    cleaned = re.sub(r"<[^>]+>", " ", text or "")
    cleaned = unescape(cleaned)
    return normalize_ws(cleaned)


def split_sentences_simple(text: str) -> List[str]:
    txt = normalize_ws(text)
    if not txt:
        return []
    parts = re.split(r"(?<=[.!?])\s+", txt)
    return [p.strip() for p in parts if p.strip()]


def keyword_terms(text: str, limit: int = 12) -> List[str]:
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9\-]{2,}", normalize_text_loose(text))
    stop = {
        "the","and","for","with","that","this","from","were","was","have","has","had","into","their","there","these","those","then","than","also","such","using","used","use","into","within","between","which","while","where","when","they","them","been","being","over","under","through","because","could","would","should","about","after","before","during","results","method","methods","paper","study","review","research","data","figure","table"
    }
    scored=[]
    seen=set()
    for tok in tokens:
        low=tok.lower()
        if low in stop or len(low)<4 or low in seen:
            continue
        seen.add(low)
        score=(2 if any(ch.isdigit() for ch in tok) else 0)+(2 if tok[0].isupper() else 0)+min(len(tok),10)/10
        scored.append((score, tok))
    scored.sort(reverse=True)
    return [tok for _,tok in scored[:limit]]


def extract_relevant_snippets(text: str, query_text: str, max_snippets: int = 5, max_chars: int = 2400) -> List[str]:
    sentences = split_sentences_simple(text)
    if not sentences:
        return []
    query_terms = {t.lower() for t in keyword_terms(query_text, limit=14)}
    ranked=[]
    for sent in sentences:
        sent_terms = {t.lower() for t in re.findall(r"[A-Za-z][A-Za-z0-9\-]{2,}", normalize_text_loose(sent))}
        overlap = len(query_terms & sent_terms)
        if overlap <= 0:
            continue
        score = overlap * 10 + min(len(sent), 240) / 120
        ranked.append((score, sent))
    ranked.sort(reverse=True)
    out=[]
    used=0
    for _, sent in ranked[:max_snippets*3]:
        if sent in out:
            continue
        if used + len(sent) > max_chars and out:
            break
        out.append(sent)
        used += len(sent)
        if len(out) >= max_snippets:
            break
    return out


def comment_priority(comment: str, kind: str, source_agent: str) -> int:
    txt = normalize_text_loose(comment).lower()
    base = {
        "fact_check_agent": 100,
        "overall_review_agent": 95,
        "literature_review_agent": 90,
        "methods_agent": 85,
        "results_agent": 85,
        "science_agent": 78,
        "crossref_reference_agent": 72,
        "citation_style_agent": 35,
        "grammar_agent": 20,
    }.get(source_agent, 50)
    for needle, bonus in [
        ("retract", 50), ("expression of concern", 40), ("correct", 20), ("corrig", 20),
        ("unsupported", 18), ("overclaim", 18), ("not clearly supported", 18), ("contradict", 18),
        ("future year", 12), ("truncated", 14), ("incomplete", 12), ("corrupt", 12), ("ocr artifact", 10),
        ("year mismatch", 15), ("mismatch", 10), ("human attention", 10)
    ]:
        if needle in txt:
            base += bonus
    return base


def is_minor_comment(comment: str, kind: str, source_agent: str) -> bool:
    txt = normalize_text_loose(comment).lower()
    if source_agent == "grammar_agent":
        return True
    minor_markers = [
        "punctuation", "capitalization", "spacing", "typo", "en dash", "hyphen", "whitespace",
        "doi prefix", "superscript", "page range", "title case", "sentence case", "style consistency",
        "extra spaces", "full-width period", "formatting artifact", "remove redundant space"
    ]
    serious_markers = [
        "retract", "expression of concern", "correct", "unsupported", "overclaim", "not clearly supported",
        "contradict", "future year", "truncated", "incomplete", "corrupt", "ocr artifact", "year mismatch"
    ]
    if any(m in txt for m in serious_markers):
        return False
    if source_agent == "citation_style_agent" and any(m in txt for m in minor_markers):
        return True
    if source_agent == "crossref_reference_agent" and "possible missing doi" in txt:
        return True
    return False


def sanitize_edit_comment(kind: str, source_agent: str, comment: str) -> str:
    if not comment:
        return ""
    if source_agent in {"grammar_agent", "citation_style_agent"}:
        return ""
    if kind in {"grammar", "citation_style", "punctuation", "concision", "style", "clarity"} and len(comment) < 120:
        return ""
    return comment


def should_surface_comment(comment: str, kind: str, source_agent: str) -> bool:
    if not normalize_ws(comment):
        return False
    if is_minor_comment(comment, kind, source_agent):
        return False
    return True


def cap_comments_per_paragraph(comments: List[CommentSuggestion], max_per_paragraph: int) -> List[CommentSuggestion]:
    grouped: Dict[Tuple[str, str], List[CommentSuggestion]] = defaultdict(list)
    for c in comments:
        key = (c.paragraph_id or f"FIG:{c.figure_id}", c.kind)
        grouped[key].append(c)
    out: List[CommentSuggestion] = []
    bucketed: Dict[str, List[CommentSuggestion]] = defaultdict(list)
    for c in comments:
        pid = c.paragraph_id or f"FIG:{c.figure_id}"
        bucketed[pid].append(c)
    for pid, items in bucketed.items():
        items.sort(key=lambda c: comment_priority(c.comment, c.kind, c.source_agent), reverse=True)
        seen=set()
        kept=0
        for c in items:
            sig = normalize_text_loose(c.comment)
            if sig in seen:
                continue
            seen.add(sig)
            out.append(c)
            kept += 1
            if kept >= max_per_paragraph:
                break
    out.sort(key=lambda c: (c.paragraph_id or c.figure_id, -comment_priority(c.comment, c.kind, c.source_agent)))
    return out


# -----------------------------------------------------------------------------
# DOCX traversal and extraction
# -----------------------------------------------------------------------------


def iter_block_paragraphs(parent: _Document | _Cell) -> Iterator[Paragraph]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_paragraphs(cell)


def detect_section(paragraph: Paragraph, current_section: str) -> str:
    style_name = getattr(paragraph.style, "name", "") or ""
    txt = normalize_ws(paragraph.text)
    if not txt:
        return current_section

    if style_name.lower().startswith("heading"):
        return txt

    heading_like = (
        len(txt) < 90
        and not txt.endswith(".")
        and txt.lower() not in {"figure", "table"}
        and re.match(r"^([A-Z][A-Za-z0-9\-–—,:() ]+|\d+(\.\d+)*\s+.+)$", txt)
    )
    if heading_like:
        return txt
    return current_section


def section_bucket(section_name: str) -> str:
    s = normalize_ws(section_name).lower()
    if not s or s == "front matter":
        return "front_matter"
    if any(k in s for k in ["abstract"]):
        return "abstract"
    if any(k in s for k in ["introduction", "background", "related work", "literature review"]):
        return "introduction"
    if any(k in s for k in ["method", "methods", "methodology", "materials", "experimental", "experiment setup"]):
        return "methods"
    if any(k in s for k in ["result", "evaluation", "findings", "analysis"]):
        return "results"
    if any(k in s for k in ["discussion", "limitations"]):
        return "discussion"
    if any(k in s for k in ["conclusion", "future work", "summary"]):
        return "conclusion"
    if any(k in s for k in ["reference", "bibliography"]):
        return "references"
    if any(k in s for k in ["appendix", "supplementary"]):
        return "appendix"
    return "other"


def looks_like_caption(text: str) -> bool:
    txt = normalize_ws(text)
    return bool(re.match(r"^(fig\.?|figure|table)\s*\d+[A-Za-z0-9\-.: ]*", txt, flags=re.I))


def sniff_binary_image_mime(path: Path) -> str:
    try:
        head = path.read_bytes()[:64]
    except Exception:
        return ""
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if head.startswith((b"GIF87a", b"GIF89a")):
        return "image/gif"
    if head.startswith(b"BM"):
        return "image/bmp"
    if head[:4] in {b"II*\x00", b"MM\x00*"}:
        return "image/tiff"
    if len(head) >= 12 and head[:4] == b"RIFF" and head[8:12] == b"WEBP":
        return "image/webp"
    return ""


def detect_best_image_mime(path: Path, declared_mime: str = "") -> str:
    sniffed = sniff_binary_image_mime(path)
    if sniffed:
        return sniffed
    declared = (declared_mime or "").lower().strip()
    if declared and declared != "application/octet-stream":
        return declared
    guessed = mimetypes.guess_type(path.name)[0] or ""
    return guessed.lower()


def extract_figures(docx_path: Path, out_dir: Path) -> List[FigureRecord]:
    figures_dir = out_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)

    records: List[FigureRecord] = []
    with zipfile.ZipFile(docx_path, "r") as zf:
        media_files = sorted(name for name in zf.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        for idx, member in enumerate(media_files, start=1):
            filename = Path(member).name
            extracted_path = figures_dir / filename
            with zf.open(member) as src, open(extracted_path, "wb") as dst:
                shutil.copyfileobj(src, dst)
            mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
            mime = detect_best_image_mime(extracted_path, mime) or mime
            records.append(
                FigureRecord(
                    figure_id=f"fig{idx:03d}",
                    filename=filename,
                    extracted_path=str(extracted_path),
                    mime_type=mime,
                    caption_guess="",
                )
            )
    return records


def build_review_payload(docx_path: Path, work_dir: Path) -> Dict[str, Any]:
    doc = Document(str(docx_path))
    paragraphs: List[ParagraphRecord] = []
    current_section = "Front Matter"

    for order, paragraph in enumerate(iter_block_paragraphs(doc), start=1):
        txt = paragraph.text or ""
        current_section = detect_section(paragraph, current_section)
        style_name = getattr(paragraph.style, "name", "Normal") or "Normal"
        bucket = section_bucket(current_section)
        paragraphs.append(
            ParagraphRecord(
                paragraph_id=f"p{order:04d}",
                order=order,
                section=current_section,
                section_bucket=bucket,
                style=style_name,
                text=txt,
                text_hash=text_hash(txt),
                in_table=isinstance(paragraph._parent, _Cell),
                contains_citation=contains_in_text_citation(txt),
            )
        )

    figures = extract_figures(docx_path, work_dir)
    caption_records = [p for p in paragraphs if looks_like_caption(p.text)]
    for fig, cap in zip(figures, caption_records):
        fig.caption_guess = cap.text
        fig.caption_paragraph_id = cap.paragraph_id

    full_text = "\n\n".join([p.text for p in paragraphs if normalize_ws(p.text)])
    references_start = None
    for i, p in enumerate(paragraphs):
        if p.section_bucket == "references" or normalize_ws(p.text).lower() in {"references", "bibliography"}:
            references_start = i
            break

    references_paragraphs = []
    if references_start is not None:
        references_paragraphs = [asdict(p) for p in paragraphs[references_start + 1 :] if normalize_ws(p.text)]

    payload = {
        "source_docx": str(docx_path),
        "paragraph_count": len(paragraphs),
        "figure_count": len(figures),
        "paragraphs": [asdict(p) for p in paragraphs],
        "figures": [asdict(f) for f in figures],
        "full_text": full_text,
        "references_detected": references_start is not None,
        "references_text": "\n".join(p["text"] for p in references_paragraphs),
        "references_paragraphs": references_paragraphs,
    }
    return payload



# -----------------------------------------------------------------------------
# Initial document analysis / classification
# -----------------------------------------------------------------------------


def summarize_document_for_prepass(payload: Dict[str, Any], max_paragraphs: int = 18) -> Dict[str, Any]:
    paragraphs = payload.get("paragraphs", [])
    nonempty = [p for p in paragraphs if normalize_ws(p.get("text", ""))]
    section_counts: Dict[str, int] = defaultdict(int)
    section_examples: Dict[str, List[str]] = defaultdict(list)

    for p in nonempty:
        bucket = p.get("section_bucket", "other")
        section_counts[bucket] += 1
        if len(section_examples[bucket]) < 3:
            section_examples[bucket].append(normalize_ws(p.get("text", ""))[:500])

    preview = []
    for p in nonempty[:max_paragraphs]:
        preview.append(
            {
                "paragraph_id": p["paragraph_id"],
                "section": p.get("section", ""),
                "section_bucket": p.get("section_bucket", "other"),
                "text": p.get("text", "")[:1200],
            }
        )

    return {
        "paragraph_count": len(nonempty),
        "figure_count": payload.get("figure_count", 0),
        "references_detected": payload.get("references_detected", False),
        "section_counts": dict(section_counts),
        "section_examples": dict(section_examples),
        "preview_paragraphs": preview,
    }


def initial_document_review(client: OpenAI, model: str, payload: Dict[str, Any], retries: int) -> Dict[str, Any]:
    heuristic_has_methods = any(p.get("section_bucket") == "methods" for p in payload.get("paragraphs", []))
    heuristic_has_results = any(p.get("section_bucket") in {"results", "discussion"} for p in payload.get("paragraphs", []))
    heuristic_doc_type = "research_paper" if (heuristic_has_methods or heuristic_has_results) else "literature_review"

    system_prompt = textwrap.dedent(
        f"""
        You are the IntakeEditor for a multi-agent PhD-paper reviewer.

        Tasks:
        1. Classify the document as one of:
           - "research_paper"
           - "literature_review"
           - "mixed_or_unclear"
        2. Decide whether the paper appears to contain its own methods/results.
        3. Give a concise overall review of structure, logical flow, and top-level scientific-writing issues.
        4. Point out whether claims appear broadly cautious or overclaimed from a high-level reading.
        5. Recommend which specialist agents should be emphasized.

        Important:
        - A literature review may have analysis/synthesis but no original experimental methods/results.
        - Do not inspect claim truth against external sources here; only assess the document itself.
        - Keep comments concrete and actionable.

        Heuristic prior:
        - has_methods={heuristic_has_methods}
        - has_results={heuristic_has_results}
        - heuristic_doc_type={heuristic_doc_type}

        Return strict JSON only:
        {{
          "document_type": "research_paper",
          "has_own_methods": true,
          "has_own_results": true,
          "reasoning": "1-3 short sentences",
          "overall_summary": "short paragraph",
          "overall_strengths": ["item"],
          "overall_risks": ["item"],
          "document_comments": [
            {{
              "paragraph_id": "p0001",
              "comment": "overall document-level comment to show in Word near the start",
              "kind": "overall_review"
            }}
          ],
          "agent_emphasis": {{
            "grammar": "high",
            "science": "high",
            "methods": "high",
            "results": "medium",
            "citation_style": "high",
            "literature_review": "low",
            "fact_check": "high"
          }}
        }}
        """
    ).strip()

    user_text = json.dumps(summarize_document_for_prepass(payload), ensure_ascii=False, indent=2)
    try:
        data = call_json_model(client, model, system_prompt, user_text, retries=retries)
    except Exception as exc:
        logger.warning("Initial document review failed; using heuristic fallback: %s", exc)
        data = {
            "document_type": heuristic_doc_type,
            "has_own_methods": heuristic_has_methods,
            "has_own_results": heuristic_has_results,
            "reasoning": "Heuristic classification fallback.",
            "overall_summary": "",
            "overall_strengths": [],
            "overall_risks": [],
            "document_comments": [],
            "agent_emphasis": {},
        }

    data.setdefault("document_type", heuristic_doc_type)
    data.setdefault("has_own_methods", heuristic_has_methods)
    data.setdefault("has_own_results", heuristic_has_results)
    data.setdefault("document_comments", [])
    data.setdefault("agent_emphasis", {})
    return data


def select_paragraphs_for_fact_check(payload: Dict[str, Any], limit: int) -> List[Dict[str, Any]]:
    candidates = []
    for p in payload.get("paragraphs", []):
        if p.get("section_bucket") == "references":
            continue
        txt = p.get("text", "")
        if not normalize_ws(txt):
            continue
        if not contains_in_text_citation(txt):
            continue
        score = len(extract_in_text_citations(txt)) * 10 + min(len(txt), 400) // 80
        candidates.append((score, p))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return [p for _, p in candidates[: max(0, limit)]]
# -----------------------------------------------------------------------------
# OpenAI / xAI client helpers
# -----------------------------------------------------------------------------


def require_api_key() -> str:
    if not DEFAULT_API_KEY:
        raise RuntimeError("No API key found. Set OPENAI_API_KEY or XAI_API_KEY in your environment or .env file.")
    return DEFAULT_API_KEY


def build_client(base_url: Optional[str], api_key: Optional[str], timeout: int, max_retries: int) -> OpenAI:
    return OpenAI(
        api_key=api_key or require_api_key(),
        base_url=base_url or DEFAULT_BASE_URL,
        timeout=timeout,
        max_retries=max_retries,
    )


def _strip_json_fences(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def _extract_first_json_object(text: str) -> str:
    text = _strip_json_fences(text)
    if text.startswith("{") and text.endswith("}"):
        return text
    match = re.search(r"\{.*\}", text, flags=re.S)
    if match:
        return match.group(0)
    raise ValueError("No JSON object found in model response")


def call_json_model(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_text: str,
    temperature: float = 0.1,
    retries: int = 3,
) -> Dict[str, Any]:
    last_exc: Optional[Exception] = None
    for attempt in range(1, retries + 1):
        try:
            response = client.responses.create(
                model=model,
                temperature=temperature,
                input=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_text},
                ],
            )
            raw = (response.output_text or "").strip()
            data = json.loads(_extract_first_json_object(raw))
            if isinstance(data, list):
                data = {"edits": data, "comments": []}
            if not isinstance(data, dict):
                raise ValueError(f"Model returned non-object JSON: {type(data).__name__}")
            return data
        except Exception as exc:  # pragma: no cover - depends on remote API
            last_exc = exc
            wait_s = min(2 ** (attempt - 1), 8)
            logger.warning("Model call failed on attempt %s/%s: %s", attempt, retries, exc)
            if attempt < retries:
                time.sleep(wait_s)
    raise RuntimeError(f"Failed to parse/use model JSON after {retries} attempts: {last_exc}")


# -----------------------------------------------------------------------------
# Vision helper
# -----------------------------------------------------------------------------


def maybe_summarize_figures_with_vision(
    client: OpenAI,
    model: str,
    payload: Dict[str, Any],
    enabled: bool,
    max_figures: int = 4,
) -> None:
    if not enabled:
        return

    supported_mimes = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif", "image/bmp", "image/tiff"}
    figures = payload.get("figures", [])[:max_figures]
    for fig in figures:
        path = Path(fig["extracted_path"])
        if not path.exists():
            continue
        mime = detect_best_image_mime(path, str(fig.get("mime_type") or "")) or ""
        fig["mime_type"] = mime or fig.get("mime_type", "")
        if mime not in supported_mimes:
            logger.info("Skipping non-image figure asset for vision: %s (%s)", path.name, mime or fig.get("mime_type") or "unknown")
            fig["vision_summary"] = ""
            continue
        b64 = base64.b64encode(path.read_bytes()).decode("ascii")
        prompt = textwrap.dedent(
            f"""
            Summarize this research-paper figure for downstream review.
            Return strict JSON only:
            {{
              "summary": "2-4 sentence neutral description",
              "possible_issues": ["optional item", "optional item"]
            }}

            Caption guess: {fig.get('caption_guess', '')}
            Focus on what the figure appears to show, not whether it is scientifically correct.
            """
        ).strip()
        try:
            response = client.responses.create(
                model=model,
                input=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": prompt},
                            {"type": "input_image", "image_url": f"data:{mime};base64,{b64}"},
                        ],
                    }
                ],
            )
            data = json.loads(_extract_first_json_object((response.output_text or "").strip()))
            summary = data.get("summary", "") if isinstance(data, dict) else ""
            issues = data.get("possible_issues", []) if isinstance(data, dict) else []
            fig["vision_summary"] = summary + (" Issues: " + "; ".join(issues) if issues else "")
        except Exception as exc:  # pragma: no cover - remote/API specific
            logger.warning("Vision summarization skipped for %s: %s", path.name, exc)
            fig["vision_summary"] = ""


# -----------------------------------------------------------------------------
# Chunking and section filters
# -----------------------------------------------------------------------------


def chunk_paragraphs(
    paragraphs: Sequence[Dict[str, Any]],
    max_chars: int,
    include_buckets: Optional[set[str]] = None,
    exclude_buckets: Optional[set[str]] = None,
    require_citations: bool = False,
    focus_paragraph_ids: Optional[set[str]] = None,
) -> List[List[Dict[str, Any]]]:
    chunks: List[List[Dict[str, Any]]] = []
    current: List[Dict[str, Any]] = []
    current_len = 0

    for p in paragraphs:
        txt = p.get("text", "")
        if not normalize_ws(txt):
            continue
        bucket = p.get("section_bucket", "other")
        if include_buckets and bucket not in include_buckets:
            continue
        if exclude_buckets and bucket in exclude_buckets:
            continue
        if require_citations and not (p.get("contains_citation") or bucket == "references"):
            continue
        if focus_paragraph_ids is not None and p["paragraph_id"] not in focus_paragraph_ids:
            continue

        serial = json.dumps(
            {
                "paragraph_id": p["paragraph_id"],
                "section": p["section"],
                "section_bucket": bucket,
                "text": txt,
            },
            ensure_ascii=False,
        )
        n = len(serial)
        if current and current_len + n > max_chars:
            chunks.append(current)
            current = []
            current_len = 0
        current.append({
            "paragraph_id": p["paragraph_id"],
            "section": p["section"],
            "section_bucket": bucket,
            "text": txt,
        })
        current_len += n

    if current:
        chunks.append(current)
    return chunks


# -----------------------------------------------------------------------------
# Agent prompts
# -----------------------------------------------------------------------------


def grammar_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], mode: str, aggressiveness: str, retries: int) -> Dict[str, Any]:
    aggressiveness_rule = {"conservative": "Prefer smaller edits unless the wording is clearly poor.", "balanced": "Use moderate edits when they clearly improve readability.", "substantive": "You may substantially rewrite a sentence or a whole paragraph locally when it clearly improves clarity, logic, and academic tone."}.get(aggressiveness, "Use moderate edits when they clearly improve readability.")
    edit_rule = "Do not propose edits; use comments only." if mode == "comments_only" else f"Make tracked-change edits. {aggressiveness_rule}"
    system_prompt = textwrap.dedent(
        f"""
        You are the GrammarAgent in a multi-agent paper-review system.

        Task:
        - Improve grammar, spelling, punctuation, concision, and academic style.
        - {edit_rule}
        - You may make sentence-level or paragraph-local rewrites when they clearly improve clarity and flow.
        - Do not change scientific meaning.
        - Do not invent citations.
        - Prefer tracked changes over comments.
        - Do not add comments for minor punctuation, capitalization, spacing, or routine grammar fixes.
        - Only add a comment when the problem likely needs human judgment, such as ambiguity, structural weakness, or unclear logic.
        - Prefer stronger edits when the wording is awkward, repetitive, or non-native, but keep each edit anchored to the source paragraph.

        Return strict JSON only:
        {{
          "edits": [
            {{
              "paragraph_id": "p0001",
              "find": "original snippet",
              "replace": "replacement snippet",
              "context_before": "up to 30 chars before",
              "context_after": "up to 30 chars after",
              "comment": "optional short rationale",
              "kind": "grammar",
              "confidence": "high"
            }}
          ],
          "comments": [
            {{
              "paragraph_id": "p0001",
              "anchor_text": "optional snippet",
              "comment": "short comment for the author",
              "kind": "grammar"
            }}
          ]
        }}

        Rules for edits:
        - Keep edits exact and applyable.
        - Only propose edits when the exact 'find' text exists in the paragraph, unless you are intentionally replacing the whole paragraph.
        - It is acceptable to rewrite a full paragraph when the paragraph is clearly weak, redundant, or illogical.
        """
    ).strip()
    return call_json_model(client, model, system_prompt, json.dumps({"paragraphs": paragraph_chunk}, ensure_ascii=False, indent=2), retries=retries)


def science_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], figures: List[Dict[str, Any]], mode: str, aggressiveness: str, retries: int) -> Dict[str, Any]:
    aggressiveness_rule = {"conservative": "Keep edits minimal and prefer comments.", "balanced": "You may revise scientifically imprecise wording when the intended meaning is clear.", "substantive": "You may rewrite sentences for scientific precision when the intended meaning is clear, but never invent facts."}.get(aggressiveness, "You may revise scientifically imprecise wording when the intended meaning is clear.")
    edit_rule = "Comments only; do not propose edits." if mode == "comments_only" else f"Use comments for risky scientific issues; {aggressiveness_rule}"
    system_prompt = textwrap.dedent(
        f"""
        You are the ScienceAgent in a multi-agent paper-review system.

        Task:
        - Review scientific clarity, internal consistency, unsupported claims, and precise academic phrasing.
        - {edit_rule}
        - You may propose stronger wording changes than a proofreader when they improve scientific precision.
        - Never claim a result is wrong unless the issue is directly visible in the text or explicitly contradicted by provided evidence.
        - Never fabricate references, numbers, or methodology details.
        - Make obvious fixes directly as edits.
        - Reserve comments for issues that genuinely need human attention, such as unsupported claims, conceptual ambiguity, weak sourcing, or substantive scientific risk.

        Return strict JSON only with keys 'edits' and 'comments'.
        Each edit must be a short local replacement and each comment must be specific.
        """
    ).strip()
    user_payload = {"paragraphs": paragraph_chunk, "figures": figures[:4]}
    return call_json_model(client, model, system_prompt, json.dumps(user_payload, ensure_ascii=False, indent=2), retries=retries)


def methods_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], retries: int) -> Dict[str, Any]:
    system_prompt = textwrap.dedent(
        """
        You are the MethodsAgent in a multi-agent paper-review system.

        Task:
        - Review methods / experimental sections for reproducibility and reporting quality.
        - Focus on missing sample sizes, unclear datasets, missing controls, unspecified hyperparameters,
          ambiguous preprocessing, undefined variables, and vague statistical methodology.
        - Use comments only.

        Return strict JSON only:
        {"edits": [], "comments": [{"paragraph_id":"p0001","anchor_text":"","comment":"specific methods review note","kind":"methods"}]}
        """
    ).strip()
    return call_json_model(client, model, system_prompt, json.dumps({"paragraphs": paragraph_chunk}, ensure_ascii=False, indent=2), retries=retries)


def results_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], retries: int) -> Dict[str, Any]:
    system_prompt = textwrap.dedent(
        """
        You are the ResultsAgent in a multi-agent paper-review system.

        Task:
        - Review results / discussion / conclusion paragraphs for overclaiming, weak evidence language,
          causal overreach, unsupported generalization, and unclear comparison statements.
        - Use comments only.

        Return strict JSON only:
        {"edits": [], "comments": [{"paragraph_id":"p0001","anchor_text":"","comment":"specific results/discussion review note","kind":"results"}]}
        """
    ).strip()
    return call_json_model(client, model, system_prompt, json.dumps({"paragraphs": paragraph_chunk}, ensure_ascii=False, indent=2), retries=retries)


def citation_style_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], references_text: str, mode: str, retries: int) -> Dict[str, Any]:
    edit_rule = "Comments only; do not propose edits." if mode == "comments_only" else "Edits are allowed only for safe local citation/reference formatting fixes."
    system_prompt = textwrap.dedent(
        f"""
        You are the CitationStyleAgent in a multi-agent paper-review system.

        Task:
        - Review in-text citations and reference list formatting for consistency.
        - {edit_rule}
        - Never fabricate bibliographic details.
        - Do not claim that an in-text citation is missing from the reference list. Reference-presence checking is handled deterministically elsewhere.
        - Use edits for punctuation, spacing, capitalization, duplicated DOI prefixes, obvious OCR cleanup, superscript artifacts, and very safe style consistency fixes.
        - Reserve comments for issues that likely need human attention, such as suspicious future years, truncated/corrupted entries, incomplete bibliographic fields, retractions/corrections, or unresolved ambiguity.
        - Do not add comments for minor punctuation, capitalization, spacing, or routine formatting fixes if they can be corrected directly.

        Return strict JSON only with keys 'edits' and 'comments'.
        """
    ).strip()
    user_payload = {"paragraphs": paragraph_chunk, "references_text": references_text[:12000]}
    return call_json_model(client, model, system_prompt, json.dumps(user_payload, ensure_ascii=False, indent=2), retries=retries)



def overall_review_agent(
    client: OpenAI,
    model: str,
    payload: Dict[str, Any],
    document_profile: Dict[str, Any],
    retries: int,
) -> Dict[str, Any]:
    system_prompt = textwrap.dedent(
        """
        You are the OverallReviewAgent in a multi-agent paper-review system.

        Task:
        - Review the whole document at a high level.
        - Focus on structure, narrative flow, section ordering, redundancy, logical transitions,
          thesis clarity, and whether conclusions are aligned with the presented material.
        - Keep feedback specific and practical.
        - Use comments only.

        Return strict JSON only:
        {
          "document_summary": "short paragraph",
          "comments": [
            {"paragraph_id": "p0001", "anchor_text": "", "comment": "specific overall comment", "kind": "overall_review"}
          ]
        }
        """
    ).strip()
    summary_payload = {
        "document_profile": document_profile,
        "document_summary_input": summarize_document_for_prepass(payload, max_paragraphs=24),
        "ending_preview": [
            {
                "paragraph_id": p["paragraph_id"],
                "section": p.get("section", ""),
                "section_bucket": p.get("section_bucket", "other"),
                "text": p.get("text", "")[:900],
            }
            for p in payload.get("paragraphs", [])[-12:]
            if normalize_ws(p.get("text", ""))
        ],
    }
    return call_json_model(client, model, system_prompt, json.dumps(summary_payload, ensure_ascii=False, indent=2), retries=retries)


def literature_review_agent(client: OpenAI, model: str, paragraph_chunk: List[Dict[str, Any]], retries: int) -> Dict[str, Any]:
    system_prompt = textwrap.dedent(
        """
        You are the LiteratureReviewAgent in a multi-agent paper-review system.

        Task:
        - Review a literature review or background-heavy section.
        - Focus on synthesis vs. summary, comparison of sources, thematic organization,
          chronology problems, missing critical evaluation, and unsupported transitions.
        - Use comments only.
        - Do not demand methods/results sections if this appears to be a literature review.

        Return strict JSON only:
        {"edits": [], "comments": [{"paragraph_id":"p0001","anchor_text":"","comment":"specific literature-review note","kind":"literature_review"}]}
        """
    ).strip()
    return call_json_model(client, model, system_prompt, json.dumps({"paragraphs": paragraph_chunk}, ensure_ascii=False, indent=2), retries=retries)


def _reconstruct_openalex_abstract(inv: Dict[str, List[int]]) -> str:
    if not isinstance(inv, dict) or not inv:
        return ""
    positions: Dict[int, str] = {}
    for token, idxs in inv.items():
        for idx in idxs or []:
            if idx not in positions:
                positions[idx] = token
    if not positions:
        return ""
    return " ".join(positions[i] for i in sorted(positions))


def openalex_lookup_by_doi(doi: str, timeout: int) -> Dict[str, Any]:
    normalized = (doi or "").strip().lower()
    if not normalized:
        return {}
    try:
        r = requests.get(
            "https://api.openalex.org/works",
            params={"filter": f"doi:https://doi.org/{normalized}", "per-page": 1},
            timeout=timeout,
            headers={"User-Agent": "profe-ai-paper-reviewer/3.0"},
        )
        if not r.ok:
            return {}
        results = (r.json() or {}).get("results") or []
        if not results:
            return {}
        item = results[0]
        abstract = _reconstruct_openalex_abstract(item.get("abstract_inverted_index") or {})
        return {
            "openalex_id": item.get("id", ""),
            "title": item.get("display_name", ""),
            "publication_year": str(item.get("publication_year", "") or ""),
            "abstract": abstract,
            "is_retracted": bool(item.get("is_retracted", False)),
            "type": item.get("type", ""),
        }
    except Exception:
        return {}


def fetch_crossref_updates_for_doi(doi: str, mailto: str, timeout: int) -> List[Dict[str, str]]:
    normalized = (doi or "").strip()
    if not normalized:
        return []
    try:
        r = requests.get(
            "https://api.crossref.org/works",
            params={"filter": f"updates:{normalized}", "rows": 5},
            headers=_crossref_headers(mailto),
            timeout=timeout,
        )
        if not r.ok:
            return []
        items = ((r.json() or {}).get("message") or {}).get("items") or []
        updates=[]
        for item in items:
            title = _best_crossref_title(item)
            doi2 = item.get("DOI", "")
            rels = []
            if isinstance(item.get("relation"), dict):
                rels = list(item.get("relation", {}).keys())
            update_type = first_nonempty([item.get("update-to", ""), ", ".join(rels), item.get("type", "")])
            updates.append({"doi": doi2, "title": title, "type": update_type})
        return updates
    except Exception:
        return []


def fetch_unpaywall_record(doi: str, email: str, timeout: int) -> Dict[str, Any]:
    normalized = (doi or "").strip()
    if not normalized or not email:
        return {}
    try:
        r = requests.get(
            f"https://api.unpaywall.org/v2/{requests.utils.quote(normalized, safe='')}",
            params={"email": email},
            timeout=timeout,
            headers={"User-Agent": f"profe-ai-paper-reviewer/3.3 ({email})"},
        )
        if not r.ok:
            return {}
        data = r.json() or {}
        best = data.get("best_oa_location") or {}
        return {
            "is_oa": bool(data.get("is_oa")),
            "oa_status": data.get("oa_status", ""),
            "best_oa_url": first_nonempty([best.get("url_for_pdf", ""), best.get("url", "")]),
            "best_oa_url_for_pdf": best.get("url_for_pdf", ""),
            "best_oa_version": best.get("version", ""),
            "best_oa_license": best.get("license", ""),
            "host_type": best.get("host_type", ""),
        }
    except Exception:
        return {}


def europepmc_lookup_by_doi(doi: str, timeout: int) -> Dict[str, Any]:
    normalized = (doi or "").strip()
    if not normalized:
        return {}
    try:
        r = requests.get(
            "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
            params={"query": f'DOI:"{normalized}"', "format": "json", "pageSize": 1, "resultType": "core"},
            timeout=timeout,
            headers={"User-Agent": "profe-ai-paper-reviewer/3.3"},
        )
        if not r.ok:
            return {}
        result_list = ((r.json() or {}).get("resultList") or {}).get("result") or []
        if not result_list:
            return {}
        item = result_list[0]
        pmcid = item.get("pmcid", "")
        abstract = strip_xml_or_html(item.get("abstractText", ""))
        fulltext = ""
        if pmcid and (str(item.get("isOpenAccess", "")).upper() == "Y" or str(item.get("hasBook", "")).upper() == "Y" or str(item.get("hasTextMinedTerms", "")).upper() == "Y" or str(item.get("inEPMC", "")).upper() == "Y"):
            try:
                fx = requests.get(
                    f"https://www.ebi.ac.uk/europepmc/webservices/rest/{pmcid}/fullTextXML",
                    params={"format": "xml"},
                    timeout=timeout,
                    headers={"User-Agent": "profe-ai-paper-reviewer/3.3"},
                )
                if fx.ok and fx.text:
                    fulltext = strip_xml_or_html(fx.text)
            except Exception:
                fulltext = ""
        return {
            "pmcid": pmcid,
            "title": item.get("title", ""),
            "abstract": abstract,
            "fulltext": fulltext,
            "journal": item.get("journalTitle", ""),
            "pubYear": str(item.get("pubYear", "") or ""),
            "source": "EuropePMC",
        }
    except Exception:
        return {}


def extract_text_from_pdf_bytes(data: bytes) -> str:
    if not data or PdfReader is None:
        return ""
    try:
        reader = PdfReader(BytesIO(data))
        parts=[]
        for page in reader.pages[:25]:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return normalize_ws("\n".join(parts))
    except Exception:
        return ""


def fetch_text_from_oa_url(url: str, timeout: int, max_chars: int) -> Tuple[str, str]:
    if not url:
        return "", ""
    try:
        r = requests.get(url, timeout=timeout, headers={"User-Agent": "profe-ai-paper-reviewer/3.3"}, stream=True)
        if not r.ok:
            return "", ""
        content_type = (r.headers.get("Content-Type") or "").lower()
        raw = r.content[:10_000_000]
        if "pdf" in content_type or url.lower().endswith(".pdf"):
            return extract_text_from_pdf_bytes(raw)[:max_chars], "pdf"
        text = strip_xml_or_html(raw.decode("utf-8", errors="ignore"))
        return text[:max_chars], "html"
    except Exception:
        return "", ""


def load_oa_evidence_for_doi(doi: str, cfg: ReviewConfig, work_dir: Path) -> Dict[str, Any]:
    if not cfg.open_access_fulltext_enabled:
        return {}
    normalized = (doi or "").strip().lower()
    if not normalized:
        return {}
    cache_path = work_dir / "oa_evidence_cache.json"
    cache = read_json_cache(cache_path)
    if normalized in cache:
        return cache[normalized]

    out: Dict[str, Any] = {}
    unpaywall = fetch_unpaywall_record(normalized, cfg.unpaywall_email or cfg.crossref_mailto, cfg.request_timeout)
    if unpaywall:
        out.update(unpaywall)
        url = first_nonempty([unpaywall.get("best_oa_url_for_pdf", ""), unpaywall.get("best_oa_url", "")])
        if url:
            text, fmt = fetch_text_from_oa_url(url, cfg.request_timeout, cfg.max_fulltext_chars)
            if text:
                out["fulltext"] = text
                out["fulltext_format"] = fmt
                out["fulltext_source"] = "Unpaywall"
                out["fulltext_url"] = url
    if not out.get("fulltext"):
        epmc = europepmc_lookup_by_doi(normalized, cfg.request_timeout)
        if epmc:
            out.setdefault("epmc_title", epmc.get("title", ""))
            out.setdefault("epmc_abstract", epmc.get("abstract", ""))
            if epmc.get("fulltext"):
                out["fulltext"] = epmc.get("fulltext", "")[: cfg.max_fulltext_chars]
                out["fulltext_format"] = "xml"
                out["fulltext_source"] = "EuropePMC"
                out["pmcid"] = epmc.get("pmcid", "")
            elif epmc.get("abstract"):
                out.setdefault("fallback_abstract", epmc.get("abstract", ""))
    cache[normalized] = out
    write_json_cache(cache_path, cache)
    return out


def build_reference_catalog(payload: Dict[str, Any], cfg: ReviewConfig, work_dir: Path) -> List[Dict[str, Any]]:
    cache_path = work_dir / "reference_lookup_cache.json"
    cache = read_json_cache(cache_path)
    ref_paragraphs = payload.get("references_paragraphs", [])[: cfg.max_references_to_verify]
    results: List[Dict[str, Any]] = []

    def _lookup(ref: Dict[str, Any]) -> Dict[str, Any]:
        ref_text = ref.get("text", "")
        key = text_hash(ref_text)
        if key in cache:
            cached = dict(cache[key])
            cached["paragraph_id"] = ref["paragraph_id"]
            cached["reference_text"] = ref_text
            return cached

        finding = crossref_lookup_reference(ref_text, cfg.crossref_mailto, timeout=cfg.request_timeout)
        doi = finding.crossref_doi or finding.found_doi
        oa = openalex_lookup_by_doi(doi, timeout=cfg.request_timeout) if doi else {}
        updates = fetch_crossref_updates_for_doi(doi, cfg.crossref_mailto, timeout=cfg.request_timeout) if doi else []
        year = extract_reference_year(ref_text)
        surname_match = re.match(r"([A-Z][A-Za-z'`\-]+)", normalize_ws(ref_text))
        entry = {
            "paragraph_id": ref["paragraph_id"],
            "reference_text": ref_text,
            "lead_author": surname_match.group(1).lower() if surname_match else "",
            "year": year[:4] if year else "",
            "found_doi": finding.found_doi,
            "crossref_doi": finding.crossref_doi,
            "crossref_title": finding.crossref_title,
            "crossref_year": finding.crossref_year,
            "matched": finding.matched,
            "issues": finding.issues,
            "lookup_method": finding.lookup_method,
            "openalex_title": oa.get("title", ""),
            "openalex_year": oa.get("publication_year", ""),
            "openalex_abstract": oa.get("abstract", ""),
            "openalex_retracted": oa.get("is_retracted", False),
            "openalex_type": oa.get("type", ""),
            "crossref_updates": updates,
        }
        cache[key] = entry
        return entry

    if cfg.reference_workers <= 1:
        for ref in ref_paragraphs:
            results.append(_lookup(ref))
    else:
        with concurrent.futures.ThreadPoolExecutor(max_workers=cfg.reference_workers) as ex:
            futures = [ex.submit(_lookup, ref) for ref in ref_paragraphs]
            for fut in concurrent.futures.as_completed(futures):
                results.append(fut.result())

    write_json_cache(cache_path, cache)
    results.sort(key=lambda x: x.get("paragraph_id", ""))
    return results


def reference_catalog_index(catalog: Sequence[Dict[str, Any]]) -> Dict[Tuple[str, str], List[Dict[str, Any]]]:
    idx: Dict[Tuple[str, str], List[Dict[str, Any]]] = defaultdict(list)
    for entry in catalog:
        sig = (entry.get("lead_author", ""), entry.get("year", ""))
        if sig[0] and sig[1]:
            idx[sig].append(entry)
    return idx


def fact_check_agent(
    client: OpenAI,
    model: str,
    paragraph: Dict[str, Any],
    cited_evidence: List[Dict[str, Any]],
    document_profile: Dict[str, Any],
    mode: str,
    retries: int,
) -> Dict[str, Any]:
    if not cited_evidence:
        return {"edits": [], "comments": []}

    edit_rule = (
        "Use comments only."
        if mode == "comments_only"
        else "Edits are allowed only for very safe claim-softening changes such as changing absolute language to cautious language."
    )
    system_prompt = textwrap.dedent(
        f"""
        You are the FactCheckAgent in a multi-agent paper-review system.

        Task:
        - Compare the paragraph's explicit scientific claims against the cited-source evidence provided.
        - Use only the supplied internet lookup evidence derived from reference metadata, abstracts, open-access full text, and post-publication status signals when available.
        - Identify overclaiming, citation mismatch, citation-status risk, or claims not clearly supported by the cited evidence.
        - {edit_rule}
        - Do not invent facts.
        - If the evidence is too thin, say so in a comment rather than asserting the claim is false.
        - If the required correction is obvious and low-risk, propose a direct tracked edit and optionally add a brief comment only if human attention is still needed.

        Return strict JSON only:
        {{
          "edits": [
            {{
              "paragraph_id": "p0001",
              "find": "exact risky wording",
              "replace": "more cautious wording",
              "comment": "why softened",
              "kind": "fact_check",
              "confidence": "medium"
            }}
          ],
          "comments": [
            {{
              "paragraph_id": "p0001",
              "anchor_text": "",
              "comment": "specific fact-check note linked to the cited evidence",
              "kind": "fact_check"
            }}
          ]
        }}
        """
    ).strip()
    user_payload = {
        "document_type": document_profile.get("document_type", ""),
        "paragraph": paragraph,
        "cited_evidence": cited_evidence[:8],
    }
    return call_json_model(client, model, system_prompt, json.dumps(user_payload, ensure_ascii=False, indent=2), retries=retries)

# -----------------------------------------------------------------------------
# Deterministic reference verification (Crossref)
# -----------------------------------------------------------------------------


def extract_doi(text: str) -> str:
    m = DOI_RE.search(text or "")
    return m.group(1).rstrip(".,;)") if m else ""


def extract_reference_year(text: str) -> str:
    m = YEAR_RE.search(text or "")
    return m.group(0) if m else ""


def extract_reference_title_guess(text: str) -> str:
    txt = normalize_ws(text)
    if not txt:
        return ""
    txt = DOI_RE.sub("", txt)

    # Rough heuristic: title often appears after year and before next period.
    year_match = YEAR_RE.search(txt)
    if year_match:
        after = txt[year_match.end():].strip(" .;:()")
        parts = re.split(r"\.\s+", after)
        if parts:
            candidate = parts[0].strip(" .;:()")
            if 10 <= len(candidate) <= 220:
                return candidate

    # Fallback: longest sentence-like segment.
    segments = [s.strip(" .;:()") for s in re.split(r"\.\s+", txt)]
    segments = [s for s in segments if 10 <= len(s) <= 220]
    segments.sort(key=len, reverse=True)
    return segments[0] if segments else ""


def _crossref_headers(mailto: str) -> Dict[str, str]:
    ua = "profe-ai-paper-reviewer/2.0"
    if mailto:
        ua += f" (mailto:{mailto})"
    return {"User-Agent": ua, "Accept": "application/json"}


def _best_crossref_title(item: Dict[str, Any]) -> str:
    titles = item.get("title") or []
    return normalize_ws(titles[0]) if titles else ""


def _best_crossref_year(item: Dict[str, Any]) -> str:
    for key in ["published-print", "published-online", "created", "issued"]:
        part = item.get(key) or {}
        date_parts = part.get("date-parts") or []
        if date_parts and date_parts[0]:
            return str(date_parts[0][0])
    return ""


def reference_similarity(a: str, b: str) -> float:
    aw = set(re.findall(r"[A-Za-z0-9]+", (a or "").lower()))
    bw = set(re.findall(r"[A-Za-z0-9]+", (b or "").lower()))
    if not aw or not bw:
        return 0.0
    return len(aw & bw) / max(1, len(aw | bw))


def crossref_lookup_reference(ref_text: str, mailto: str, timeout: int) -> ReferenceCheckFinding:
    doi = extract_doi(ref_text)
    title_guess = extract_reference_title_guess(ref_text)
    ref_year = extract_reference_year(ref_text)
    finding = ReferenceCheckFinding(
        paragraph_id="",
        reference_text=ref_text,
        found_doi=doi,
        confidence="low",
    )

    session = requests.Session()
    headers = _crossref_headers(mailto)

    try:
        if doi:
            encoded = requests.utils.requote_uri(doi)
            url = f"https://api.crossref.org/works/{encoded}"
            r = session.get(url, headers=headers, timeout=timeout)
            if r.ok:
                item = (r.json() or {}).get("message") or {}
                finding.crossref_doi = item.get("DOI", "")
                finding.crossref_title = _best_crossref_title(item)
                finding.crossref_year = _best_crossref_year(item)
                finding.matched = True
                finding.lookup_method = "doi"
                finding.confidence = "high"
                sim = reference_similarity(title_guess, finding.crossref_title)
                if title_guess and finding.crossref_title and sim < 0.35:
                    finding.issues.append("DOI resolves, but the resolved title looks different from the reference text.")
                if ref_year and finding.crossref_year and ref_year[:4] != finding.crossref_year[:4]:
                    finding.issues.append(f"Year mismatch: reference says {ref_year[:4]}, Crossref says {finding.crossref_year[:4]}.")
                return finding
            finding.issues.append(f"DOI lookup failed with HTTP {r.status_code}.")

        query_text = ref_text[:500]
        r = session.get(
            "https://api.crossref.org/works",
            params={"query.bibliographic": query_text, "rows": 1},
            headers=headers,
            timeout=timeout,
        )
        if not r.ok:
            finding.issues.append(f"Crossref bibliographic lookup failed with HTTP {r.status_code}.")
            return finding

        items = ((r.json() or {}).get("message") or {}).get("items") or []
        if not items:
            finding.issues.append("No Crossref match found.")
            return finding

        item = items[0]
        finding.crossref_doi = item.get("DOI", "")
        finding.crossref_title = _best_crossref_title(item)
        finding.crossref_year = _best_crossref_year(item)
        finding.lookup_method = "bibliographic"
        sim = reference_similarity(title_guess, finding.crossref_title or ref_text)
        finding.matched = sim >= 0.35 or bool(finding.crossref_doi)
        finding.confidence = "medium" if finding.matched else "low"

        if title_guess and finding.crossref_title and sim < 0.35:
            finding.issues.append("Low-confidence title match from Crossref.")
        if not doi and finding.crossref_doi:
            finding.issues.append(f"Possible missing DOI: {finding.crossref_doi}")
        if ref_year and finding.crossref_year and ref_year[:4] != finding.crossref_year[:4]:
            finding.issues.append(f"Year mismatch: reference says {ref_year[:4]}, Crossref says {finding.crossref_year[:4]}.")
        if not finding.issues and finding.matched:
            finding.issues.append("Reference appears to match Crossref.")
        return finding
    except Exception as exc:  # pragma: no cover - network specific
        finding.issues.append(f"Crossref lookup error: {exc}")
        return finding


def extract_reference_signatures(paragraphs: Sequence[Dict[str, Any]]) -> set[Tuple[str, str]]:
    signatures: set[Tuple[str, str]] = set()
    for p in paragraphs:
        if p.get("section_bucket") != "references":
            continue
        txt = normalize_text_loose(p.get("text", ""))
        if not txt:
            continue
        year = extract_reference_year(txt)
        author_match = re.match(r"(.+?)\(\s*(19|20)\d{2}[a-z]?\s*\)", txt)
        if not year or not author_match:
            continue
        surnames = split_reference_author_block(author_match.group(1))
        if not surnames:
            continue
        for surname in surnames[:3]:
            signatures.add((surname, year[:4]))
    return signatures


def extract_in_text_citations(text: str) -> List[Tuple[str, str]]:
    found: List[Tuple[str, str]] = []
    txt = normalize_text_loose(text or "")
    for pattern in CITATION_PATTERNS:
        for m in pattern.finditer(txt):
            author_part = m.group(1)
            year_match = re.search(r"(19|20)\d{2}", m.group(0))
            if not year_match:
                continue
            surnames = split_reference_author_block(author_part)
            if surnames:
                found.append((surnames[0], year_match.group(0)))
    return found


def verify_references_and_citations(
    payload: Dict[str, Any],
    cfg: ReviewConfig,
    work_dir: Path,
) -> Tuple[List[CommentSuggestion], List[ReferenceCheckFinding], List[Dict[str, Any]]]:
    comments: List[CommentSuggestion] = []
    findings: List[ReferenceCheckFinding] = []

    catalog = build_reference_catalog(payload, cfg, work_dir) if cfg.crossref_enabled else []

    for entry in catalog:
        finding = ReferenceCheckFinding(
            paragraph_id=entry.get("paragraph_id", ""),
            reference_text=entry.get("reference_text", ""),
            found_doi=entry.get("found_doi", ""),
            crossref_doi=entry.get("crossref_doi", ""),
            crossref_title=entry.get("crossref_title", "") or entry.get("openalex_title", ""),
            crossref_year=entry.get("crossref_year", "") or entry.get("openalex_year", ""),
            confidence="high" if entry.get("matched") else "low",
            issues=list(entry.get("issues") or []),
            matched=bool(entry.get("matched")),
            lookup_method=entry.get("lookup_method", ""),
        )
        if entry.get("openalex_retracted"):
            finding.issues.append("OpenAlex flags this work as retracted.")
        findings.append(finding)

        actionable = [x for x in finding.issues if x and not x.startswith("Reference appears to match") and "Possible missing DOI" not in x]
        serious_update_types = ", ".join(sorted({u.get("type", "") for u in entry.get("crossref_updates", []) if u.get("type")}))
        if serious_update_types:
            actionable.append(f"Post-publication update signal detected: {serious_update_types}.")
        if actionable:
            comments.append(
                CommentSuggestion(
                    paragraph_id=finding.paragraph_id,
                    anchor_text="",
                    comment="Reference verification: " + " ".join(actionable[:3]),
                    kind="reference_verification",
                    source_agent="crossref_reference_agent",
                )
            )

    ref_signatures = extract_reference_signatures(payload.get("paragraphs", []))
    for p in payload.get("paragraphs", []):
        if p.get("section_bucket") == "references":
            continue
        para_citations = extract_in_text_citations(p.get("text", ""))
        missing = []
        for sig in para_citations:
            if sig not in ref_signatures:
                missing.append(f"{sig[0].title()} {sig[1]}")
        if missing:
            uniq = sorted(set(missing))
            comments.append(
                CommentSuggestion(
                    paragraph_id=p["paragraph_id"],
                    anchor_text="",
                    comment="Possible citation/reference mismatch: could not confidently match in-text citation(s) to the reference list: " + ", ".join(uniq[:6]),
                    kind="citation_consistency",
                    source_agent="crossref_reference_agent",
                )
            )
    return comments, findings, catalog


# -----------------------------------------------------------------------------
# Agent execution and chief editor
# -----------------------------------------------------------------------------


@dataclass
class AgentTask:
    agent_name: str
    chunk_index: int
    round_index: int
    paragraphs: List[Dict[str, Any]]


class ChiefEditor:
    def __init__(self, cfg: ReviewConfig, api_key: Optional[str]) -> None:
        self.cfg = cfg
        self.api_key = api_key or require_api_key()
        self.reference_signatures: set[Tuple[str, str]] = set()

    def _client(self) -> OpenAI:
        return build_client(self.cfg.base_url, self.api_key, self.cfg.request_timeout, self.cfg.max_retries)

    def _run_task(self, task: AgentTask, figures: List[Dict[str, Any]], references_text: str, document_profile: Dict[str, Any]) -> Dict[str, Any]:
        client = self._client()
        logger.info(
            "Round %s | %s | chunk %s | paragraphs=%s",
            task.round_index,
            task.agent_name,
            task.chunk_index + 1,
            len(task.paragraphs),
        )
        if task.agent_name == "grammar_agent":
            result = grammar_agent(client, self.cfg.model, task.paragraphs, self.cfg.mode, self.cfg.editing_aggressiveness, self.cfg.max_retries)
        elif task.agent_name == "science_agent":
            result = science_agent(client, self.cfg.model, task.paragraphs, figures, self.cfg.mode, self.cfg.editing_aggressiveness, self.cfg.max_retries)
        elif task.agent_name == "methods_agent":
            result = methods_agent(client, self.cfg.model, task.paragraphs, self.cfg.max_retries)
        elif task.agent_name == "results_agent":
            result = results_agent(client, self.cfg.model, task.paragraphs, self.cfg.max_retries)
        elif task.agent_name == "citation_style_agent":
            result = citation_style_agent(client, self.cfg.model, task.paragraphs, references_text, self.cfg.mode, self.cfg.max_retries)
        elif task.agent_name == "literature_review_agent":
            result = literature_review_agent(client, self.cfg.model, task.paragraphs, self.cfg.max_retries)
        else:
            raise ValueError(f"Unknown agent: {task.agent_name}")
        result["_meta"] = {"agent_name": task.agent_name, "chunk_index": task.chunk_index, "round_index": task.round_index}
        return result

    def _build_tasks(self, payload: Dict[str, Any], focus_paragraph_ids: Optional[set[str]], round_index: int) -> List[AgentTask]:
        paragraphs = payload["paragraphs"]
        tasks: List[AgentTask] = []
        document_profile = payload.get("document_profile", {})
        doc_type = document_profile.get("document_type", "research_paper")

        specs: List[Tuple[str, List[List[Dict[str, Any]]]]] = []

        if doc_type == "literature_review":
            specs.append(("literature_review_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, include_buckets={"abstract", "introduction", "discussion", "conclusion", "other"}, focus_paragraph_ids=focus_paragraph_ids)))
        else:
            specs.extend([
                ("methods_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, include_buckets={"methods"}, focus_paragraph_ids=focus_paragraph_ids)),
                ("results_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, include_buckets={"results", "discussion", "conclusion"}, focus_paragraph_ids=focus_paragraph_ids)),
            ])

        specs.extend([
            ("science_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, include_buckets={"abstract", "introduction", "results", "discussion", "conclusion", "other"}, focus_paragraph_ids=focus_paragraph_ids)),
            ("citation_style_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, include_buckets={"references", "abstract", "introduction", "methods", "results", "discussion", "conclusion", "other"}, require_citations=False, focus_paragraph_ids=focus_paragraph_ids)),
            ("grammar_agent", chunk_paragraphs(paragraphs, self.cfg.chunk_chars, exclude_buckets={"references"}, focus_paragraph_ids=focus_paragraph_ids)),
        ])

        for agent_name, chunks in specs:
            for idx, chunk in enumerate(chunks):
                if not chunk:
                    continue
                tasks.append(AgentTask(agent_name=agent_name, chunk_index=idx, round_index=round_index, paragraphs=chunk))
        return tasks

    def _coerce_result_item(self, item: Any, *, item_type: str, agent_name: str) -> Optional[Dict[str, Any]]:
        if isinstance(item, dict):
            data = dict(item)
        elif isinstance(item, str):
            raw = item.strip()
            if not raw:
                return None
            # Try strict JSON first.
            try:
                parsed = json.loads(raw)
                if isinstance(parsed, dict):
                    data = dict(parsed)
                else:
                    logger.warning("Skipping non-object %s from %s: %r", item_type, agent_name, parsed)
                    return None
            except Exception:
                if item_type == "comment":
                    parsed_comment = self._parse_plain_comment_string(raw, agent_name)
                    if parsed_comment:
                        data = parsed_comment
                    else:
                        logger.warning("Skipping non-JSON %s from %s: %r", item_type, agent_name, raw[:200])
                        return None
                else:
                    logger.warning("Skipping non-JSON %s from %s: %r", item_type, agent_name, raw[:200])
                    return None
        elif isinstance(item, (list, tuple)):
            try:
                data = dict(item)
            except Exception:
                logger.warning("Skipping malformed %s from %s: %r", item_type, agent_name, item)
                return None
        else:
            logger.warning("Skipping unsupported %s type from %s: %s", item_type, agent_name, type(item).__name__)
            return None

        if item_type == "edit":
            data = self._normalize_edit_dict(data)
        if item_type == "comment":
            if "comment" not in data and "text" in data:
                data["comment"] = data.pop("text")
            data = self._normalize_comment_dict(data)
        return data

    def _normalize_edit_dict(self, data: Dict[str, Any]) -> Dict[str, Any]:
        data = dict(data)
        if "find" not in data and "old" in data:
            data["find"] = data.pop("old")
        if "replace" not in data and "new" in data:
            data["replace"] = data.pop("new")
        if "find" not in data and "from" in data:
            data["find"] = data.pop("from")
        if "replace" not in data and "to" in data:
            data["replace"] = data.pop("to")
        if "find" not in data and "original_snippet" in data:
            data["find"] = data.pop("original_snippet")
        if "replace" not in data and "replacement_snippet" in data:
            data["replace"] = data.pop("replacement_snippet")
        if "find" not in data and "original" in data:
            data["find"] = data.pop("original")
        if isinstance(data.get("replacement"), dict):
            rep = data.get("replacement") or {}
            data.setdefault("find", rep.get("old", ""))
            data.setdefault("replace", rep.get("new", ""))
        if "change" in data and not data.get("find") and not data.get("replace"):
            m = re.search(r"change\s+'(.+?)'\s+to\s+'(.+?)'", str(data.get("change", "")), flags=re.I)
            if m:
                data["find"], data["replace"] = m.group(1), m.group(2)
        replacement_value = first_nonempty([data.get("replace"), data.get("replacement_text"), data.get("new_text"), data.get("text")])
        if not data.get("find") and replacement_value and data.get("replace_entire_paragraph"):
            data["replace"] = replacement_value
        elif replacement_value and "replace" not in data:
            data["replace"] = replacement_value
        allowed = {"paragraph_id", "find", "replace", "context_before", "context_after", "comment", "kind", "confidence", "source_agent", "round_index", "replace_entire_paragraph"}
        return {k: v for k, v in data.items() if k in allowed}

    def _normalize_comment_dict(self, data: Dict[str, Any]) -> Dict[str, Any]:
        allowed = {"paragraph_id", "anchor_text", "comment", "kind", "source_agent", "round_index", "figure_id"}
        if "comment" in data:
            data["comment"] = normalize_ws(str(data.get("comment", "")))
        return {k: v for k, v in data.items() if k in allowed}

    def _parse_plain_comment_string(self, raw_text: str, agent_name: str) -> Optional[Dict[str, Any]]:
        txt = normalize_ws(raw_text)
        if not txt or len(txt) < 3:
            return None

        paragraph_id = ""
        figure_id = ""

        both = re.match(r"^(?:paragraph\s+)?(p\d{4})\s*/\s*(?:figure\s+)?(fig\d{3})(?:\s*\([^)]*\))?\s*:\s*(.+)$", txt, flags=re.I)
        if both:
            paragraph_id = both.group(1).lower()
            figure_id = both.group(2).lower()
            txt = both.group(3)
        else:
            fig_first = re.match(r"^(?:figure\s+)?(fig\d{3})(?:\s*\([^)]*\))?\s*:\s*(.+)$", txt, flags=re.I)
            if fig_first:
                figure_id = fig_first.group(1).lower()
                txt = fig_first.group(2)
            else:
                m = re.match(r"^(p\d{4})(?:\s*,\s*p\d{4})*\s*:\s*(.+)$", txt, flags=re.I)
                if m:
                    paragraph_id = m.group(1).lower()
                    txt = m.group(2)
                else:
                    m2 = re.match(r"^In\s+paragraph\s+(p\d{4})\s*:\s*(.+)$", txt, flags=re.I)
                    if m2:
                        paragraph_id = m2.group(1).lower()
                        txt = m2.group(2)
                    else:
                        fig_embedded = re.match(r"^(?:paragraph\s+)?(p\d{4})\s*.*?\b(fig\d{3})\b.*?:\s*(.+)$", txt, flags=re.I)
                        if fig_embedded:
                            paragraph_id = fig_embedded.group(1).lower()
                            figure_id = fig_embedded.group(2).lower()
                            txt = fig_embedded.group(3)

        cleaned = normalize_ws(txt)
        if not cleaned:
            return None
        return {
            "paragraph_id": paragraph_id,
            "figure_id": figure_id,
            "anchor_text": "",
            "comment": cleaned,
            "kind": agent_name.replace("_agent", ""),
        }

    def _should_drop_missing_reference_comment(self, item: Dict[str, Any], agent_name: str) -> bool:
        if agent_name != "citation_style_agent":
            return False
        comment_text = normalize_text_loose(item.get("comment", ""))
        if "not found in reference list" not in comment_text and "missing reference" not in comment_text and "lacks a corresponding entry" not in comment_text:
            return False
        cited = extract_comment_citation_signatures(comment_text)
        if not cited:
            return False
        return all(sig in self.reference_signatures for sig in cited)

    def _add_result_items(
        self,
        result: Dict[str, Any],
        merged_edits: List[EditSuggestion],
        merged_comments: List[CommentSuggestion],
    ) -> None:
        meta = result.get("_meta", {}) if isinstance(result, dict) else {}
        agent_name = meta.get("agent_name", "")
        round_index = int(meta.get("round_index", 1))

        edits = result.get("edits", []) if isinstance(result, dict) else []
        comments = result.get("comments", []) if isinstance(result, dict) else []
        if isinstance(edits, dict):
            edits = [edits]
        elif isinstance(edits, str):
            edits = [edits]
        if isinstance(comments, dict):
            comments = [comments]
        elif isinstance(comments, str):
            comments = [comments]

        for raw_item in edits or []:
            item = self._coerce_result_item(raw_item, item_type="edit", agent_name=agent_name)
            if not item:
                continue
            item.setdefault("source_agent", agent_name)
            item.setdefault("kind", agent_name.replace("_agent", ""))
            item.setdefault("round_index", round_index)
            item["comment"] = sanitize_edit_comment(item.get("kind", ""), agent_name, str(item.get("comment", "")))
            try:
                merged_edits.append(EditSuggestion(**item))
            except Exception:
                logger.warning("Skipping malformed edit from %s: %s", agent_name, item)

        for raw_item in comments or []:
            item = self._coerce_result_item(raw_item, item_type="comment", agent_name=agent_name)
            if not item and isinstance(raw_item, str):
                item = self._parse_plain_comment_string(raw_item, agent_name)
            if not item:
                continue
            item.setdefault("source_agent", agent_name)
            item.setdefault("kind", agent_name.replace("_agent", ""))
            item.setdefault("round_index", round_index)
            if self._should_drop_missing_reference_comment(item, agent_name):
                logger.info("Dropping false missing-reference comment from %s: %s", agent_name, item.get("comment", ""))
                continue
            if not should_surface_comment(str(item.get("comment", "")), str(item.get("kind", "")), agent_name):
                continue
            try:
                merged_comments.append(CommentSuggestion(**item))
            except Exception:
                logger.warning("Skipping malformed comment from %s: %s", agent_name, item)

    def _resolve_conflicts_with_chief(
        self,
        payload: Dict[str, Any],
        edits: List[EditSuggestion],
        comments: List[CommentSuggestion],
    ) -> Tuple[List[EditSuggestion], List[CommentSuggestion]]:
        by_paragraph_edits: Dict[str, List[EditSuggestion]] = defaultdict(list)
        by_paragraph_comments: Dict[str, List[CommentSuggestion]] = defaultdict(list)
        para_text = {p["paragraph_id"]: p["text"] for p in payload["paragraphs"]}

        for e in edits:
            by_paragraph_edits[e.paragraph_id].append(e)
        for c in comments:
            by_paragraph_comments[c.paragraph_id].append(c)

        accepted_edits: List[EditSuggestion] = []
        accepted_comments: List[CommentSuggestion] = []

        client = self._client()
        for paragraph_id in sorted(set(by_paragraph_edits) | set(by_paragraph_comments)):
            p_text = para_text.get(paragraph_id, "")
            p_edits = by_paragraph_edits.get(paragraph_id, [])
            p_comments = by_paragraph_comments.get(paragraph_id, [])

            conflict = self._paragraph_has_conflict(p_edits)
            heavy = len(p_edits) > 4 or len(p_comments) > 4
            if not conflict and not heavy:
                accepted_edits.extend(p_edits)
                accepted_comments.extend(p_comments)
                continue

            logger.info("ChiefEditor resolving paragraph %s (%s edits, %s comments)", paragraph_id, len(p_edits), len(p_comments))
            system_prompt = textwrap.dedent(
                f"""
                You are the ChiefEditor in a multi-agent paper-review system.
                Decide which proposed paragraph-level edits/comments to keep.
                Review mode is: {self.cfg.mode}
                Rules:
                - Preserve scientific meaning.
                - Prefer fewer, high-confidence edits.
                - Reject conflicting or redundant edits.
                - In comments_only mode, return zero edits.
                - Keep comments specific and non-duplicative.

                Return strict JSON only:
                {{
                  "keep_edit_indices": [0, 2],
                  "keep_comment_indices": [1],
                  "notes": "optional"
                }}
                """
            ).strip()
            user_payload = {
                "paragraph_id": paragraph_id,
                "paragraph_text": p_text,
                "edits": [asdict(x) for x in p_edits],
                "comments": [asdict(x) for x in p_comments],
            }
            try:
                decision = call_json_model(client, self.cfg.model, system_prompt, json.dumps(user_payload, ensure_ascii=False, indent=2), retries=self.cfg.max_retries)
                keep_edit_indices = set(decision.get("keep_edit_indices", []))
                keep_comment_indices = set(decision.get("keep_comment_indices", []))
                accepted_edits.extend([e for idx, e in enumerate(p_edits) if idx in keep_edit_indices])
                accepted_comments.extend([c for idx, c in enumerate(p_comments) if idx in keep_comment_indices])
            except Exception as exc:
                logger.warning("ChiefEditor resolution failed for %s, falling back to heuristic keep-first: %s", paragraph_id, exc)
                accepted_edits.extend(self._heuristic_keep_edits(p_edits))
                accepted_comments.extend(self._heuristic_keep_comments(p_comments))

        return accepted_edits, accepted_comments

    @staticmethod
    def _heuristic_keep_edits(edits: List[EditSuggestion]) -> List[EditSuggestion]:
        kept: List[EditSuggestion] = []
        spans: List[Tuple[str, str]] = []
        for e in edits:
            key = (e.find, e.replace)
            if key in spans:
                continue
            spans.append(key)
            kept.append(e)
        return kept[:4]

    @staticmethod
    def _heuristic_keep_comments(comments: List[CommentSuggestion]) -> List[CommentSuggestion]:
        seen = set()
        kept: List[CommentSuggestion] = []
        for c in comments:
            if c.comment in seen:
                continue
            seen.add(c.comment)
            kept.append(c)
        return kept[:4]

    @staticmethod
    def _paragraph_has_conflict(edits: List[EditSuggestion]) -> bool:
        seen_find: Dict[str, str] = {}
        for e in edits:
            if e.find in seen_find and seen_find[e.find] != e.replace:
                return True
            seen_find[e.find] = e.replace
        return False

    def run(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        figures = payload.get("figures", [])
        references_text = payload.get("references_text", "")
        document_profile = payload.get("document_profile", {})
        self.reference_signatures = extract_reference_signatures(payload.get("paragraphs", []))
        merged_edits: List[EditSuggestion] = []
        merged_comments: List[CommentSuggestion] = []
        focus_paragraph_ids: Optional[set[str]] = None
        per_round_stats: List[Dict[str, Any]] = []
        overall_review: Dict[str, Any] = {}

        if self.cfg.structure_review_enabled:
            try:
                overall_review = overall_review_agent(self._client(), self.cfg.model, payload, document_profile, self.cfg.max_retries)
                self._add_result_items(
                    {"comments": overall_review.get("comments", []), "_meta": {"agent_name": "overall_review_agent", "chunk_index": 0, "round_index": 0}},
                    merged_edits,
                    merged_comments,
                )
            except Exception as exc:
                logger.warning("OverallReviewAgent failed: %s", exc)

        for round_index in range(1, max(1, self.cfg.rounds) + 1):
            tasks = self._build_tasks(payload, focus_paragraph_ids, round_index)
            logger.info("Prepared %s agent tasks for round %s", len(tasks), round_index)
            round_edits_before = len(merged_edits)
            round_comments_before = len(merged_comments)

            if self.cfg.max_workers <= 1:
                for task in tasks:
                    result = self._run_task(task, figures, references_text, document_profile)
                    self._add_result_items(result, merged_edits, merged_comments)
            else:
                with concurrent.futures.ThreadPoolExecutor(max_workers=self.cfg.max_workers) as ex:
                    futures = [ex.submit(self._run_task, task, figures, references_text, document_profile) for task in tasks]
                    for fut in concurrent.futures.as_completed(futures):
                        result = fut.result()
                        self._add_result_items(result, merged_edits, merged_comments)

            merged_edits = dedupe_edits(merged_edits)
            merged_comments = dedupe_comments(merged_comments)
            per_round_stats.append(
                {
                    "round_index": round_index,
                    "tasks": len(tasks),
                    "new_edits": len(merged_edits) - round_edits_before,
                    "new_comments": len(merged_comments) - round_comments_before,
                }
            )

            touched = {e.paragraph_id for e in merged_edits if e.round_index == round_index}
            touched |= {c.paragraph_id for c in merged_comments if c.round_index == round_index and c.paragraph_id}
            focus_paragraph_ids = touched or focus_paragraph_ids

        resolved_edits, resolved_comments = self._resolve_conflicts_with_chief(payload, merged_edits, merged_comments)
        final_edits = dedupe_edits(resolved_edits)
        final_comments = cap_comments_per_paragraph(dedupe_comments(resolved_comments), self.cfg.max_comments_per_paragraph)
        return {
            "edits": [asdict(e) for e in final_edits],
            "comments": [asdict(c) for c in final_comments],
            "overall_review": overall_review,
            "orchestration": {"rounds": per_round_stats, "max_workers": self.cfg.max_workers, "document_type": document_profile.get("document_type", "")},
        }


def dedupe_edits(edits: List[EditSuggestion]) -> List[EditSuggestion]:
    seen = set()
    out: List[EditSuggestion] = []
    for e in edits:
        key = (e.paragraph_id, e.find, e.replace, e.kind, e.replace_entire_paragraph)
        if (not e.find and not e.replace_entire_paragraph) or e.find == e.replace:
            continue
        if key in seen:
            continue
        seen.add(key)
        out.append(e)
    return out


def dedupe_comments(comments: List[CommentSuggestion]) -> List[CommentSuggestion]:
    seen = set()
    out: List[CommentSuggestion] = []
    for c in comments:
        if not c.comment:
            continue
        key = (c.paragraph_id, c.figure_id, c.comment, c.kind)
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
    return out


# -----------------------------------------------------------------------------
# Edit application helpers
# -----------------------------------------------------------------------------


def build_live_paragraph_map(doc: Document) -> Dict[str, Paragraph]:
    mapping: Dict[str, Paragraph] = {}
    for order, paragraph in enumerate(iter_block_paragraphs(doc), start=1):
        mapping[f"p{order:04d}"] = paragraph
    return mapping


def find_all_occurrences(text: str, needle: str) -> List[int]:
    positions: List[int] = []
    if not needle:
        return positions
    start = 0
    while True:
        idx = text.find(needle, start)
        if idx == -1:
            break
        positions.append(idx)
        start = idx + len(needle)
    return positions


def locate_edit_span(paragraph_text: str, find_text: str, context_before: str = "", context_after: str = "") -> Optional[Tuple[int, int]]:
    occurrences = find_all_occurrences(paragraph_text, find_text)
    if not occurrences:
        return None
    if len(occurrences) == 1 and not context_before and not context_after:
        start = occurrences[0]
        return start, start + len(find_text)

    candidates: List[Tuple[int, int, int]] = []
    for start in occurrences:
        end = start + len(find_text)
        score = 0
        if context_before:
            left = paragraph_text[max(0, start - len(context_before)):start]
            if left.endswith(context_before):
                score += len(context_before)
        if context_after:
            right = paragraph_text[end:end + len(context_after)]
            if right.startswith(context_after):
                score += len(context_after)
        candidates.append((score, start, end))

    candidates.sort(reverse=True)
    best_score, start, end = candidates[0]
    if best_score == 0 and len(occurrences) > 1:
        return None
    return start, end


def add_bubble_comment(document: Document, paragraph: Paragraph, comment_text: str, author: str) -> bool:
    comment_text = normalize_ws(comment_text)
    if not comment_text:
        return False
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run("")
        runs = list(paragraph.runs)
    try:
        document.add_comment(runs=runs, text=comment_text, author=author, initials=safe_initials(author))
        return True
    except Exception:
        return False


EDITABLE_KINDS = {"grammar", "citation_style", "reference", "reference_verification", "science", "fact_check"}


def apply_review_to_docx(input_docx: Path, output_docx: Path, payload: Dict[str, Any], decisions: Dict[str, Any], author: str, mode: str) -> Dict[str, Any]:
    rdoc = RevisionDocument(str(input_docx))
    live_doc = rdoc.document
    paragraph_map = build_live_paragraph_map(live_doc)
    original_text_by_id = {p["paragraph_id"]: p["text"] for p in payload["paragraphs"]}
    figure_anchor_map = {fig.get("figure_id", ""): fig.get("caption_paragraph_id", "") for fig in payload.get("figures", [])}

    edits_by_paragraph: Dict[str, List[EditSuggestion]] = defaultdict(list)
    for item in decisions.get("edits", []):
        try:
            e = EditSuggestion(**item)
        except Exception:
            continue
        if mode == "comments_only":
            continue
        if e.kind not in EDITABLE_KINDS and e.source_agent not in {"grammar_agent", "citation_style_agent", "science_agent", "fact_check_agent"}:
            continue
        edits_by_paragraph[e.paragraph_id].append(e)

    applied_edits = 0
    skipped_edits = []

    for paragraph_id, edits in edits_by_paragraph.items():
        paragraph = paragraph_map.get(paragraph_id)
        original_text = original_text_by_id.get(paragraph_id, "")
        if paragraph is None or original_text is None:
            skipped_edits.extend([{"paragraph_id": paragraph_id, "reason": "paragraph missing", **asdict(e)} for e in edits])
            continue

        located: List[Tuple[int, int, EditSuggestion]] = []
        used_ranges: List[Tuple[int, int]] = []
        for e in edits:
            if e.replace_entire_paragraph:
                start, end = 0, len(original_text)
            else:
                span = locate_edit_span(original_text, e.find, e.context_before, e.context_after)
                if span is None:
                    skipped_edits.append({"paragraph_id": paragraph_id, "reason": "could not locate span", **asdict(e)})
                    continue
                start, end = span
            overlap = any(not (end <= s or start >= e2) for s, e2 in used_ranges)
            if overlap:
                skipped_edits.append({"paragraph_id": paragraph_id, "reason": "overlapping edit", **asdict(e)})
                continue
            used_ranges.append((start, end))
            located.append((start, end, e))

        located.sort(key=lambda t: t[0], reverse=True)
        rp = RevisionParagraph.from_paragraph(paragraph)
        for start, end, e in located:
            try:
                rp.replace_tracked_at(start, end, e.replace, author=author, comment=e.comment or None)
                applied_edits += 1
            except Exception as exc:
                skipped_edits.append({"paragraph_id": paragraph_id, "reason": f"replace_tracked_at failed: {exc}", **asdict(e)})

    applied_comments = 0
    skipped_comments = []
    for item in decisions.get("comments", []):
        try:
            comment = CommentSuggestion(**item)
        except Exception:
            continue
        target_id = comment.paragraph_id or figure_anchor_map.get(comment.figure_id, "")
        paragraph = paragraph_map.get(target_id)
        if paragraph is None:
            skipped_comments.append({"paragraph_id": comment.paragraph_id, "figure_id": comment.figure_id, "reason": "paragraph missing", **asdict(comment)})
            continue
        ok = add_bubble_comment(live_doc, paragraph, comment.comment, author)
        if ok:
            applied_comments += 1
        else:
            skipped_comments.append({"reason": "add_comment failed", **asdict(comment)})

    rdoc.save(str(output_docx))
    return {
        "applied_edits": applied_edits,
        "applied_comments": applied_comments,
        "skipped_edits": skipped_edits,
        "skipped_comments": skipped_comments,
    }




def run_fact_check_reviews(
    payload: Dict[str, Any],
    decisions: Dict[str, Any],
    catalog: List[Dict[str, Any]],
    cfg: ReviewConfig,
    api_key: str,
) -> Dict[str, Any]:
    if not cfg.fact_check_enabled or "fact_check_agent" not in set(cfg.enabled_agents or default_enabled_agents()):
        return {"comments": [], "edits": [], "checked_paragraphs": 0}

    document_profile = payload.get("document_profile", {})
    idx = reference_catalog_index(catalog)
    paragraphs = select_paragraphs_for_fact_check(payload, cfg.max_fact_check_paragraphs)
    if not paragraphs:
        return {"comments": [], "edits": [], "checked_paragraphs": 0}

    client = build_client(cfg.base_url, api_key, cfg.request_timeout, cfg.max_retries)
    out_comments: List[CommentSuggestion] = []
    out_edits: List[EditSuggestion] = []

    for p in paragraphs:
        citations = extract_in_text_citations(p.get("text", ""))
        cited_evidence: List[Dict[str, Any]] = []
        seen = set()
        for sig in citations:
            for entry in idx.get(sig, []):
                key = (entry.get("crossref_doi") or entry.get("found_doi") or entry.get("paragraph_id"))
                if key in seen:
                    continue
                seen.add(key)
                doi = first_nonempty([entry.get("crossref_doi", ""), entry.get("found_doi", "")])
                oa_evidence = load_oa_evidence_for_doi(doi, cfg, Path(cfg.work_dir)) if doi else {}
                reference_text = first_nonempty([oa_evidence.get("fulltext", ""), oa_evidence.get("fallback_abstract", ""), entry.get("openalex_abstract", "")])
                snippets = extract_relevant_snippets(reference_text, p.get("text", ""), max_snippets=cfg.max_evidence_snippets, max_chars=min(cfg.max_fulltext_chars, 2600))
                update_types = [u.get("type", "") for u in entry.get("crossref_updates", []) if u.get("type")]
                cited_evidence.append(
                    {
                        "lead_author": entry.get("lead_author", ""),
                        "year": entry.get("year", ""),
                        "title": first_nonempty([entry.get("openalex_title", ""), entry.get("crossref_title", "")]),
                        "doi": doi,
                        "abstract": first_nonempty([oa_evidence.get("fallback_abstract", ""), entry.get("openalex_abstract", "")])[:2500],
                        "issues": entry.get("issues", []),
                        "is_retracted": entry.get("openalex_retracted", False),
                        "update_signals": update_types,
                        "oa_source": first_nonempty([oa_evidence.get("fulltext_source", ""), oa_evidence.get("source", ""), ""]),
                        "oa_url": oa_evidence.get("fulltext_url", oa_evidence.get("best_oa_url", "")),
                        "fulltext_snippets": snippets,
                        "fulltext_available": bool(oa_evidence.get("fulltext")),
                    }
                )

        if not cited_evidence:
            continue

        try:
            from phd_reviewer.agents.fact_check_agent import run as fact_check_agent_run
            result = fact_check_agent_run(client, cfg.model, p, cited_evidence, document_profile, cfg.mode, cfg.max_retries)
        except Exception as exc:
            logger.warning("FactCheckAgent failed for %s: %s", p.get("paragraph_id"), exc)
            continue

        meta = {"agent_name": "fact_check_agent", "chunk_index": 0, "round_index": 0}
        temp = ChiefEditor(cfg, api_key)
        edits: List[EditSuggestion] = []
        comments: List[CommentSuggestion] = []
        temp._add_result_items({**result, "_meta": meta}, edits, comments)
        out_edits.extend(edits)
        out_comments.extend(comments)

    return {
        "checked_paragraphs": len(paragraphs),
        "edits": [asdict(x) for x in dedupe_edits(out_edits)],
        "comments": [asdict(x) for x in dedupe_comments(out_comments)],
    }


# -----------------------------------------------------------------------------
# CLI and config merging
# -----------------------------------------------------------------------------


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Production multi-agent DOCX reviewer for research papers.")
    parser.add_argument("--config", help="Optional YAML or JSON config file")
    parser.add_argument("--input", help="Path to input .docx file")
    parser.add_argument("--output", help="Path to output reviewed .docx file")
    parser.add_argument("--payload-json", help="Where to save extracted structured payload JSON")
    parser.add_argument("--decisions-json", help="Where to save merged agent decisions JSON")
    parser.add_argument("--apply-report-json", help="Where to save edit-application report JSON")
    parser.add_argument("--crossref-report-json", help="Where to save deterministic reference verification report JSON")
    parser.add_argument("--model", help="Model name")
    parser.add_argument("--author", help="Author/reviewer name shown in Word tracked changes/comments")
    parser.add_argument("--base-url", help="Optional OpenAI-compatible base URL")
    parser.add_argument("--rounds", type=int, help="Number of review rounds")
    parser.add_argument("--agents", help="Comma-separated list of enabled agents, e.g. grammar_agent,science_agent,citation_style_agent")
    parser.add_argument("--work-dir", help="Temporary work directory")
    parser.add_argument("--vision", action="store_true", help="Try to summarize extracted figures with a vision-capable model")
    parser.add_argument("--mode", choices=["edit_and_comment", "comments_only"], help="Whether to apply tracked edits and comments, or comments only")
    parser.add_argument("--chunk-chars", type=int, help="Approximate chunk size sent to the model")
    parser.add_argument("--max-workers", type=int, help="Parallel worker count for chunk processing")
    parser.add_argument("--request-timeout", type=int, help="Per-request timeout in seconds")
    parser.add_argument("--max-retries", type=int, help="Retry count for model/API calls")
    parser.add_argument("--crossref-mailto", help="Optional email for Crossref polite pool")
    parser.add_argument("--disable-crossref", action="store_true", help="Disable deterministic Crossref / DOI verification")
    parser.add_argument("--max-references-to-verify", type=int, help="Cap on number of reference paragraphs to verify")
    parser.add_argument("--max-fact-check-paragraphs", type=int, help="Cap on number of body paragraphs to fact-check against cited literature")
    parser.add_argument("--editing-aggressiveness", choices=["conservative", "balanced", "substantive"], help="How willing the agents are to rewrite locally")
    parser.add_argument("--disable-structure-review", action="store_true", help="Disable the initial overall structure/flow review")
    parser.add_argument("--disable-fact-check", action="store_true", help="Disable claim checking against cited literature metadata/abstracts")
    parser.add_argument("--reference-workers", type=int, help="Parallel worker count for Crossref/OpenAlex lookups")
    parser.add_argument("--unpaywall-email", help="Email used for Unpaywall open-access lookups")
    parser.add_argument("--disable-oa-fulltext", action="store_true", help="Disable open-access full-text retrieval for fact checking")
    parser.add_argument("--log-file", help="Path to log file")
    parser.add_argument("--log-level", help="Logging level, e.g. INFO or DEBUG")
    return parser.parse_args()


CONFIG_KEYS = {
    "input", "output", "payload_json", "decisions_json", "apply_report_json", "crossref_report_json",
    "model", "author", "base_url", "rounds", "enabled_agents", "work_dir", "vision", "mode", "chunk_chars", "max_workers",
    "request_timeout", "max_retries", "crossref_enabled", "crossref_mailto", "max_references_to_verify",
    "max_fact_check_paragraphs", "editing_aggressiveness", "structure_review_enabled", "fact_check_enabled",
    "reference_workers", "unpaywall_email", "open_access_fulltext_enabled", "max_oa_lookups", "max_fulltext_chars", "max_evidence_snippets", "max_comments_per_paragraph", "log_file", "log_level"
}


def merge_config(args: argparse.Namespace) -> ReviewConfig:
    file_cfg = try_read_json_or_yaml(getattr(args, "config", None))
    file_cfg = {k: v for k, v in file_cfg.items() if k in CONFIG_KEYS}

    merged: Dict[str, Any] = {
        "input": None,
        "output": None,
        **file_cfg,
    }

    cli_map = {
        "input": args.input,
        "output": args.output,
        "payload_json": args.payload_json,
        "decisions_json": args.decisions_json,
        "apply_report_json": args.apply_report_json,
        "crossref_report_json": args.crossref_report_json,
        "model": args.model,
        "author": args.author,
        "base_url": args.base_url,
        "rounds": args.rounds,
        "enabled_agents": [x.strip() for x in args.agents.split(",") if x.strip()] if args.agents else None,
        "work_dir": args.work_dir,
        "mode": args.mode,
        "chunk_chars": args.chunk_chars,
        "max_workers": args.max_workers,
        "request_timeout": args.request_timeout,
        "max_retries": args.max_retries,
        "crossref_mailto": args.crossref_mailto,
        "max_references_to_verify": args.max_references_to_verify,
        "max_fact_check_paragraphs": args.max_fact_check_paragraphs,
        "editing_aggressiveness": args.editing_aggressiveness,
        "reference_workers": args.reference_workers,
        "unpaywall_email": args.unpaywall_email,
        "log_file": args.log_file,
        "log_level": args.log_level,
    }

    for k, v in cli_map.items():
        if v is not None:
            merged[k] = v

    if args.vision:
        merged["vision"] = True
    if args.disable_crossref:
        merged["crossref_enabled"] = False
    if args.disable_structure_review:
        merged["structure_review_enabled"] = False
    if args.disable_fact_check:
        merged["fact_check_enabled"] = False
    if args.disable_oa_fulltext:
        merged["open_access_fulltext_enabled"] = False

    if not merged.get("input") or not merged.get("output"):
        raise ValueError("Both --input and --output are required, either on the command line or in the config file.")

    defaults = ReviewConfig(input=str(merged["input"]), output=str(merged["output"]))
    default_map = asdict(defaults)
    default_map.update({k: v for k, v in merged.items() if v is not None})
    return ReviewConfig(**default_map)


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------


def run_review(cfg: ReviewConfig) -> Dict[str, Any]:
    setup_logging(cfg.log_file, cfg.log_level)

    input_docx = Path(cfg.input)
    output_docx = Path(cfg.output)
    payload_json = Path(cfg.payload_json)
    decisions_json = Path(cfg.decisions_json)
    apply_report_json = Path(cfg.apply_report_json)
    crossref_report_json = Path(cfg.crossref_report_json)
    work_dir = Path(cfg.work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)

    if not input_docx.exists():
        raise FileNotFoundError(f"Input file not found: {input_docx}")
    if input_docx.suffix.lower() != ".docx":
        raise ValueError("This script currently supports .docx only, not legacy .doc")

    api_key = require_api_key()
    client = build_client(cfg.base_url, api_key, cfg.request_timeout, cfg.max_retries)

    logger.info("[1/7] Extracting structured payload from %s ...", input_docx)
    payload = build_review_payload(input_docx, work_dir)
    maybe_summarize_figures_with_vision(client, cfg.model, payload, enabled=cfg.vision)

    logger.info("[2/7] Running intake review / document classification ...")
    document_profile = initial_document_review(client, cfg.model, payload, cfg.max_retries)
    payload["document_profile"] = document_profile
    payload_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("Saved payload JSON -> %s", payload_json)
    logger.info("Detected document type: %s | own methods=%s | own results=%s", document_profile.get("document_type", "unknown"), document_profile.get("has_own_methods"), document_profile.get("has_own_results"))

    if "reference_verification" in set(cfg.enabled_agents or default_enabled_agents()):
        logger.info("[3/7] Running deterministic reference verification ...")
        ref_comments, ref_findings, reference_catalog = verify_references_and_citations(payload, cfg, work_dir)
    else:
        logger.info("[3/7] Skipping deterministic reference verification (disabled)")
        ref_comments, ref_findings, reference_catalog = [], [], []

    logger.info("[4/7] Running multi-agent review with model %s ...", cfg.model)
    chief = ChiefEditor(cfg=cfg, api_key=api_key)
    decisions = chief.run(payload)
    for c in ref_comments:
        decisions.setdefault("comments", []).append(asdict(c))

    logger.info("[5/7] Running fact-check agent against cited literature metadata, open-access full text, and status signals ...")
    fact_results = run_fact_check_reviews(payload, decisions, reference_catalog, cfg, api_key)
    decisions.setdefault("edits", []).extend(fact_results.get("edits", []))
    decisions.setdefault("comments", []).extend(fact_results.get("comments", []))

    decisions["edits"] = [asdict(e) if isinstance(e, EditSuggestion) else e for e in dedupe_edits([EditSuggestion(**x) if not isinstance(x, EditSuggestion) else x for x in decisions.get("edits", [])])]
    decisions["comments"] = [asdict(c) if isinstance(c, CommentSuggestion) else c for c in cap_comments_per_paragraph(dedupe_comments([CommentSuggestion(**x) if not isinstance(x, CommentSuggestion) else x for x in decisions.get("comments", [])]), cfg.max_comments_per_paragraph)]
    decisions["document_profile"] = document_profile
    decisions["fact_check"] = {"checked_paragraphs": fact_results.get("checked_paragraphs", 0)}
    decisions_json.write_text(json.dumps(decisions, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("Saved merged decisions -> %s", decisions_json)
    logger.info("Proposed edits: %s", len(decisions.get("edits", [])))
    logger.info("Proposed comments: %s", len(decisions.get("comments", [])))

    crossref_report_json.write_text(json.dumps([asdict(x) for x in ref_findings], ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("Saved Crossref verification report -> %s", crossref_report_json)

    logger.info("[6/7] Applying tracked changes and comments to original DOCX ...")
    apply_report = apply_review_to_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        payload=payload,
        decisions=decisions,
        author=cfg.author,
        mode=cfg.mode,
    )
    apply_report_json.write_text(json.dumps(apply_report, ensure_ascii=False, indent=2), encoding="utf-8")

    logger.info("[7/7] Saved reviewed document -> %s", output_docx)
    logger.info("Applied tracked edits: %s", apply_report["applied_edits"])
    logger.info("Applied comments: %s", apply_report["applied_comments"])
    if apply_report["skipped_edits"]:
        logger.info("Skipped edits: %s (see %s)", len(apply_report["skipped_edits"]), apply_report_json)
    if apply_report["skipped_comments"]:
        logger.info("Skipped comments: %s (see %s)", len(apply_report["skipped_comments"]), apply_report_json)

    logger.info("Done.")
    return {
        "payload": payload,
        "decisions": decisions,
        "apply_report": apply_report,
        "reference_findings": [asdict(x) for x in ref_findings],
        "document_profile": document_profile,
        "output_docx": str(output_docx),
        "log_file": str(Path(cfg.log_file)),
    }


def main() -> None:
    args = parse_args()
    cfg = merge_config(args)
    run_review(cfg)


if __name__ == "__main__":
    main()
